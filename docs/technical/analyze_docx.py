import os
import sys
from docx import Document

def analyze_document(file_path):
    try:
        doc = Document(file_path)
        print(f"\n\n==== ANALYZING: {os.path.basename(file_path)} ====\n")

        # Extract paragraphs
        print("=== DOCUMENT STRUCTURE ===")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():  # Only print non-empty paragraphs
                style = para.style.name if para.style else "Default"
                print(f"[{style}] {para.text[:100]}{'...' if len(para.text) > 100 else ''}")
        
        # Extract tables
        print("\n=== TABLES ===")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells]
                print(f"  Row {row_idx+1}: {' | '.join(row_text)}")

    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")

def main():
    directory = "/mnt/d/Projects/ai-arbeidsdeskundige_claude/Arbeidsdeskundige data"
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            full_path = os.path.join(directory, filename)
            analyze_document(full_path)

if __name__ == "__main__":
    main()