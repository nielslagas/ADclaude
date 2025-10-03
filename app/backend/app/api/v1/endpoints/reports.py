from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse
from typing import List, Dict, Any, Optional
from uuid import UUID
from datetime import datetime
import os
import json
import logging
import re

from app.core.security import verify_token
from app.models.report import Report, ReportCreate, ReportRead, ReportSectionGenerate
from app.models.ad_report_structure import ADReport, ADReportGenerator
from app.utils.ad_report_template import get_enhanced_ad_template
from app.db.database_service import db_service
from app.core.config import settings

# Import the Celery task - we'll use the string name to avoid circular imports
from app.celery_worker import celery
from app.tasks.generate_report_tasks.structured_rag_pipeline import generate_structured_content_for_section
from app.tasks.generate_report_tasks.ad_report_task import generate_enhanced_ad_report, get_available_ad_templates

router = APIRouter()
logger = logging.getLogger(__name__)


def get_ordered_sections(template_id: str = "staatvandienst"):
    """
    Get sections in professional order based on template configuration
    
    Returns:
        List of section IDs in proper professional order
    """
    template = MVP_TEMPLATES.get(template_id, MVP_TEMPLATES["staatvandienst"])
    sections = template["sections"]
    
    # Sort by order field, fallback to original order if not specified
    ordered_sections = sorted(
        sections.items(),
        key=lambda x: x[1].get("order", 999)
    )
    
    return [section_id for section_id, _ in ordered_sections]


def get_section_mapping():
    """
    Map current/old section names to new professional section names
    This helps migrate existing reports to the new structure
    """
    return {
        # Current sections -> New professional sections
        "persoonsgegevens": "gegevensverzameling_werknemer",
        "werkgever_functie": "gegevensverzameling_werkgever", 
        "aanleiding": "vraagstelling",
        "arbeidsverleden": "gegevensverzameling_werknemer",
        "medische_situatie": "gegevensverzameling_voorgeschiedenis",
        "belastbaarheid": "belastbaarheid",
        "belasting_huidige_functie": "eigen_functie",
        "visie_ad": "visie_ad_eigen_werk",
        "matching": "visie_ad_ander_werk_extern",
        "conclusie": "conclusie",
        "samenvatting": "conclusie"
    }


def standardize_report_structure(report_content: dict, template_id: str = "staatvandienst"):
    """
    Standardize report content to professional structure
    
    Args:
        report_content: Current report content dictionary
        template_id: Template to use for standardization
        
    Returns:
        Standardized report content with professional section order
    """
    if not report_content:
        return report_content
        
    section_mapping = get_section_mapping()
    ordered_sections = get_ordered_sections(template_id)
    standardized_content = {}
    
    # Map existing sections to new structure
    for old_section_id, content in report_content.items():
        new_section_id = section_mapping.get(old_section_id, old_section_id)
        
        # If the new section already has content, combine it
        if new_section_id in standardized_content:
            standardized_content[new_section_id] += f"\n\n{content}"
        else:
            standardized_content[new_section_id] = content
    
    # Ensure all professional sections exist (even if empty)
    for section_id in ordered_sections:
        if section_id not in standardized_content:
            standardized_content[section_id] = ""
    
    return standardized_content

# Enhanced AD template - alle rapporten gebruiken nu deze template
ENHANCED_AD_TEMPLATE = {
    "enhanced_ad_rapport": {
        "id": "enhanced_ad_rapport",
        "name": "Enhanced AD Rapport",
        "description": "Professioneel arbeidsdeskundig rapport met complete structuur",
        "sections": {}  # Sections worden dynamisch gegenereerd
    }
}

# Legacy MVP_TEMPLATES voor backwards compatibility - redirect naar Enhanced AD
MVP_TEMPLATES = {
    "staatvandienst": {
        "id": "staatvandienst",
        "name": "Staatvandienst Format",
        "description": "Professioneel rapportage format volgens Nederlandse AD-standaarden",
        "sections": {
            "vraagstelling": {
                "title": "1. Vraagstelling",
                "description": "Onderzoeksvragen conform arbeidsdeskundig standaardprotocol",
                "order": 1
            },
            "ondernomen_activiteiten": {
                "title": "2. Ondernomen activiteiten", 
                "description": "Overzicht van uitgevoerde onderzoekshandelingen",
                "order": 2
            },
            "gegevensverzameling_voorgeschiedenis": {
                "title": "3.1 Voorgeschiedenis",
                "description": "Medische voorgeschiedenis en verzuimhistorie",
                "order": 3
            },
            "gegevensverzameling_werkgever": {
                "title": "3.2 Gegevens werkgever",
                "description": "Bedrijfsgegevens, aard bedrijf, functies en organisatie",
                "order": 4
            },
            "gegevensverzameling_werknemer": {
                "title": "3.3 Gegevens werknemer",
                "description": "Opleidingen, arbeidsverleden en bekwaamheden",
                "order": 5
            },
            "belastbaarheid": {
                "title": "3.4 Belastbaarheid van werknemer",
                "description": "Functionele Mogelijkhedenlijst (FML) en capaciteitsanalyse",
                "order": 6
            },
            "eigen_functie": {
                "title": "3.5 Eigen functie werknemer",
                "description": "Functieomschrijving en belastende aspecten huidige functie",
                "order": 7
            },
            "gesprek_werkgever": {
                "title": "3.6 Gesprek met de werkgever",
                "description": "Visie werkgever op functioneren en re-integratiemogelijkheden",
                "order": 8
            },
            "gesprek_werknemer": {
                "title": "3.7 Gesprek met werknemer",
                "description": "Visie werknemer op beperkingen, mogelijkheden en re-integratie",
                "order": 9
            },
            "gesprek_gezamenlijk": {
                "title": "3.8 Gesprek met werkgever en werknemer gezamenlijk",
                "description": "Afstemming en bespreking van advies",
                "order": 10
            },
            "visie_ad_eigen_werk": {
                "title": "4.1 Geschiktheid voor eigen werk",
                "description": "Arbeidsdeskundige analyse geschiktheid huidige functie",
                "order": 11
            },
            "visie_ad_aanpassing": {
                "title": "4.2 Aanpassing eigen werk", 
                "description": "Mogelijkheden voor aanpassing huidige functie",
                "order": 12
            },
            "visie_ad_ander_werk_eigen": {
                "title": "4.3 Geschiktheid voor ander werk bij eigen werkgever",
                "description": "Alternatieven binnen huidige organisatie",
                "order": 13
            },
            "visie_ad_ander_werk_extern": {
                "title": "4.4 Geschiktheid voor ander werk bij andere werkgever", 
                "description": "Zoekrichting en mogelijkheden externe arbeidsmarkt",
                "order": 14
            },
            "visie_ad_duurzaamheid": {
                "title": "4.5 Duurzaamheid van herplaatsing",
                "description": "Prognose en duurzaamheidsanalyse re-integratie",
                "order": 15
            },
            "advies": {
                "title": "5. Advies",
                "description": "Concrete aanbevelingen en vervolgstappen",
                "order": 16
            },
            "conclusie": {
                "title": "6. Conclusie",
                "description": "Beantwoording hoofdvragen en eindconclusies",
                "order": 17
            },
            "vervolg": {
                "title": "7. Vervolg",
                "description": "Vervolgtraject en aanbevolen acties",
                "order": 18
            }
        }
    },
    # Enhanced AD template based on comprehensive professional analysis
    **{get_enhanced_ad_template()["id"]: get_enhanced_ad_template()}
}

# Layout styling options for reports (not templates, but styling choices)
LAYOUT_OPTIONS = {
    "standaard": {
        "id": "standaard",
        "name": "Standaard Layout",
        "description": "Eenvoudige, klassieke opmaak"
    },
    "modern": {
        "id": "modern", 
        "name": "Modern Layout",
        "description": "Modern design met koptekst/voettekst"
    },
    "professioneel": {
        "id": "professioneel",
        "name": "Professioneel Layout", 
        "description": "Minimalistisch en zakelijk"
    }
}

@router.get("/templates", response_model=Dict[str, Any])
async def get_report_templates(user_info = Depends(verify_token)):
    """
    Get available report templates
    """
    try:
        print("Fetching templates - User info:", user_info)
        # For MVP, return the hardcoded templates
        # In a full implementation, these would be fetched from the database
        response = MVP_TEMPLATES
        print(f"Templates found: {list(response.keys())}")
        return response
    except Exception as e:
        print(f"Error fetching templates: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        # Return empty templates rather than failing
        return {}

@router.get("/layouts", response_model=Dict[str, Any])
async def get_layout_options():
    """
    Get available layout styling options
    """
    return LAYOUT_OPTIONS

@router.post("/", status_code=status.HTTP_201_CREATED, deprecated=True)
async def create_report(report_data: ReportCreate, use_structured_output: bool = False, user_info = Depends(verify_token)):
    """
    DEPRECATED: Redirects to Enhanced AD workflow
    Use /reports/enhanced-ad endpoint directly instead
    """
    logger.warning("Deprecated create_report endpoint called, redirecting to Enhanced AD")
    user_id = user_info["user_id"]

    # Debug the input
    print(f"Creating report with data: {report_data}")
    print(f"Report data type: {type(report_data)}")

    # Extract data from the request, excluding layout_type
    report_dict = report_data.dict()
    print(f"Report data dict: {report_dict}")

    # Set a default value for layout_type if not provided
    layout_type = report_dict.pop("layout_type", "standaard") if hasattr(report_data, "layout_type") else "standaard"
    print(f"Using layout_type: {layout_type} (stored in metadata, not as a direct field)")

    # Check if template exists
    if report_data.template_id not in MVP_TEMPLATES:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Template not found"
        )

    # Check if case exists and belongs to user
    try:
        case = db_service.get_case(str(report_data.case_id), user_id)

        if not case:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Case not found"
            )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error checking case: {str(e)}"
        )
    
    # Create report record
    try:
        # We already have the layout_type from earlier, but check if template wants to override it
        # If template has a specific layout and we're using the default layout, use template layout
        if layout_type == "standaard" and "layout" in MVP_TEMPLATES.get(report_data.template_id, {}):
            layout_type = MVP_TEMPLATES[report_data.template_id].get("layout", "standaard")

        # Database service will handle storing layout_type and structured metadata
        report = db_service.create_report(
            case_id=str(report_data.case_id),
            user_id=user_id,
            title=report_data.title,
            template_id=report_data.template_id,
            layout_type=layout_type,
            generation_method="structured" if use_structured_output else "legacy",
            use_structured_output=use_structured_output
        )
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to create report"
            )
        
        # Use Celery to generate the report asynchronously
        report_id = report["id"]
        
        try:
            # Set report status to processing
            db_service.update_report(report_id, {
                "status": "processing",
                "updated_at": datetime.utcnow().isoformat()
            })
            
            # Log debugging info
            print(f"About to send Celery task for report {report_id}")
            print(f"Using structured output: {use_structured_output}")
            
            # Get user profile
            try:
                user_profile = db_service.get_user_profile(user_id)
                print(f"Got user profile for user {user_id}: {user_profile is not None}")
            except Exception as profile_err:
                print(f"Error getting user profile, proceeding without it: {str(profile_err)}")
                user_profile = None

            # Always use Enhanced AD workflow - old workflow is deprecated
            logger.info(f"Starting Enhanced AD report generation for report {report_id}")
            
            # Convert old template_id to enhanced_ad_rapport if needed
            if report_data.template_id != "enhanced_ad_rapport":
                logger.info(f"Converting template {report_data.template_id} to enhanced_ad_rapport")
                db_service.update_report(report_id, {
                    "template_id": "enhanced_ad_rapport"
                })
            
            # Use Enhanced AD report task
            task = celery.send_task(
                "app.tasks.generate_report_tasks.ad_report_task.generate_enhanced_ad_report",
                args=[str(report_id)]
            )
            
            logger.info(f"Enhanced AD report task started with ID: {task.id}")
        except Exception as e:
            # Update report with error if we can't even start the task
            db_service.update_report(report_id, {
                "status": "failed",
                "error": f"Failed to start report generation: {str(e)}",
                "updated_at": datetime.utcnow().isoformat()
            })
            print(f"Error starting report generation: {str(e)}")
        
        # Add debug logging about report structure
        print(f"Report created successfully: {report}")
        print(f"Report type: {type(report)}")
        if isinstance(report, dict):
            print(f"Report keys: {report.keys()}")

        # Return the report object (with status=processing)
        # The frontend will poll for updates
        return report
    except Exception as e:
        import traceback
        print(f"Error creating report: {str(e)}")
        print(f"Error type: {type(e)}")
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error creating report: {str(e)}"
        )

@router.get("/case/{case_id}", response_model=List[ReportRead])
async def list_case_reports(case_id: UUID, user_info = Depends(verify_token)):
    """
    List all reports for a specific case
    """
    user_id = user_info["user_id"]
    
    try:
        # Check if case exists and belongs to user
        case = db_service.get_case(str(case_id), user_id)
        
        if not case:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Case not found"
            )
            
        # Get reports for the case
        reports = db_service.get_reports_for_case(str(case_id))
        
        return reports
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error listing reports: {str(e)}"
        )

# Enhanced AD Report Endpoints (moved to avoid route conflicts)
@router.get("/ad-templates", response_model=List[Dict[str, Any]])
async def get_enhanced_ad_templates(user_info = Depends(verify_token)):
    """
    Get available enhanced AD report templates based on professional analysis
    """
    try:
        templates = get_available_ad_templates()
        return templates
    except Exception as e:
        logger.error(f"Error fetching enhanced AD templates: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error fetching enhanced AD templates: {str(e)}"
        )

@router.post("/enhanced-ad", status_code=status.HTTP_201_CREATED)
async def create_enhanced_ad_report(
    report_data: ReportCreate, 
    user_info = Depends(verify_token)
):
    """
    Create a new enhanced AD report using comprehensive professional analysis
    
    This endpoint creates reports using the enhanced template structure identified
    through analysis of 4 professional AD reports, incorporating:
    - Standardized 7-section framework
    - 4-question assessment protocol
    - Detailed FML with 6 rubrieken
    - Professional formatting requirements
    """
    user_id = user_info["user_id"]
    
    try:
        print(f"Creating enhanced AD report - User: {user_id}")
        
        # Validate case exists and belongs to user
        case = db_service.get_case(str(report_data.case_id), user_id)
        if not case:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Case not found or access denied"
            )
        
        # Check for processed documents (both "processed" and "enhanced" status)
        processed_docs = db_service.get_rows("document", {
            "case_id": str(report_data.case_id),
            "status": "processed"
        })
        enhanced_docs = db_service.get_rows("document", {
            "case_id": str(report_data.case_id), 
            "status": "enhanced"
        })
        documents = processed_docs + enhanced_docs
        
        if not documents:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No processed documents found for this case. Please upload and process documents first."
            )
        
        # Validate template (force enhanced template)
        template_id = "enhanced_ad_rapport"
        enhanced_template = get_enhanced_ad_template()
        
        print(f"Using enhanced AD template with {len(enhanced_template['sections'])} sections")
        
        # Create report entry in database
        report_id = db_service.create_report(
            case_id=str(report_data.case_id),
            user_id=user_id,
            title=report_data.title,
            template_id=template_id,
            layout_type=report_data.layout_type or "standaard",
            generation_method="enhanced_ad",
            use_structured_output=True
        )
        
        # Extract the actual report ID from the returned object
        actual_report_id = report_id["id"] if isinstance(report_id, dict) else report_id
        print(f"Created enhanced AD report entry: {actual_report_id}")
        
        # Back to enhanced AD report generation with faster Haiku model
        task_result = generate_enhanced_ad_report.delay(str(actual_report_id))
        
        print(f"Triggered enhanced AD report generation task: {task_result.id}")
        
        return {
            "message": "Enhanced AD report generation started",
            "report_id": actual_report_id,
            "task_id": task_result.id,
            "template_used": template_id,
            "template_version": enhanced_template["version"],
            "sections_count": len(enhanced_template["sections"]),
            "based_on_analysis": enhanced_template["based_on_analysis"],
            "quality_checkpoints": len(enhanced_template["quality_checkpoints"])
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating enhanced AD report: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error creating enhanced AD report: {str(e)}"
        )

@router.get("/{report_id}", response_model=ReportRead)
async def get_report(report_id: UUID, user_info = Depends(verify_token)):
    """
    Get a specific report by ID
    """
    user_id = user_info["user_id"]
    
    try:
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Standardize report structure to professional format
        if "content" in report and isinstance(report["content"], dict):
            # Check if we have structured_data (from new AD report generator)
            if "structured_data" not in report["content"] and report["content"]:
                # This is legacy content, process as before
                # Get the template ID from report or use default
                template_id = report.get("template_id", "staatvandienst")
                
                # Get ordered sections for the template
                ordered_sections = get_ordered_sections(template_id)
                
                # Create new ordered content dictionary
                ordered_content = {}
                for section_id in ordered_sections:
                    if section_id in report["content"]:
                        # Clean markdown formatting from content
                        content = report["content"][section_id]
                        if isinstance(content, str):
                            import re
                            # Remove markdown headers at start of lines
                            content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
                            # Remove bold markdown (but preserve text inside)
                            content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
                            # Remove italic markdown (single asterisk, but preserve bullet points)
                            content = re.sub(r'(?<!\n)\*([^*\n]+)\*', r'\1', content)
                            # Remove underscore bold/italic
                            content = re.sub(r'__([^_]+)__', r'\1', content)
                            content = re.sub(r'_([^_]+)_', r'\1', content)
                        ordered_content[section_id] = content
                
                # Update report with ordered content
                report["content"] = ordered_content
            # If we have structured_data, keep it as is for the frontend to use
            
        return report
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error retrieving report: {str(e)}"
        )

@router.post("/regenerate-section", response_model=Dict[str, Any])
async def regenerate_report_section(
    section_data: ReportSectionGenerate, 
    user_info = Depends(verify_token)
):
    """
    Regenerate a specific section of a report
    """
    user_id = user_info["user_id"]
    
    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(section_data.report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Check if section exists in the template
        template_id = report["template_id"]
        if template_id not in MVP_TEMPLATES:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )
            
        template = MVP_TEMPLATES[template_id]
        if section_data.section_id not in template["sections"]:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Section '{section_data.section_id}' not found in template"
            )
        
        # Import at function level to avoid circular import
        # Get section variables ready for both approaches
        section_id = section_data.section_id
        section_info = template["sections"][section_id]
        
        # Direct approach implementation
        try:
            # Import needed libraries
            import google.generativeai as genai
            from app.core.config import settings
            genai.configure(api_key=settings.GOOGLE_API_KEY)
            
            # Initialize model
            model = genai.GenerativeModel(
                model_name='gemini-1.5-pro',
                generation_config={
                    "temperature": 0.1,
                    "top_p": 0.95, 
                    "top_k": 40,
                    "max_output_tokens": 4096,
                    "candidate_count": 1
                }
            )
            
            # System instruction
            system_instruction = (
                "Je bent een ervaren arbeidsdeskundige die professionele rapporten opstelt "
                "volgens de Nederlandse standaarden. Gebruik formele, zakelijke taal en "
                "zorg voor een objectieve, feitelijke weergave. "
                "Zoek zeer grondig in alle documenten naar specifieke informatie zoals persoonsgegevens, "
                "werkgever, functie, arbeidsverleden, opleiding, en medische informatie. "
                "Besteed extra aandacht aan het extraheren van feitelijke gegevens. "
                "Als je bepaalde informatie niet kunt vinden, geef dan duidelijk aan dat "
                "deze 'niet is aangetroffen in de aangeleverde documenten'."
            )
            
            # Get document chunks for this case (using report's case_id)
            case_id = report["case_id"]
            documents = db_service.get_rows("document", {
                "case_id": case_id,
                "status": ["processed", "enhanced"]
            })
            
            # Get full document content
            all_content = ""
            for doc in documents:
                # Get document chunks
                chunks = db_service.get_document_chunks(doc["id"])
                
                # Combine chunk content
                doc_content = "\n\n".join([chunk["content"] for chunk in chunks])
                all_content += f"\n\n=== DOCUMENT: {doc['filename']} ===\n\n{doc_content}"
            
            # Create sectie-specifieke prompt op basis van het sectietype
            base_prompt = f"""Schrijf een professionele sectie '{section_info['title']}' voor een arbeidsdeskundig rapport.
            Gebaseerd op de volgende documentatie, maak een zakelijke en objectieve analyse.
            Wees bondig maar volledig, en gebruik een professionele toon.
            """
            
            section_specific_instructions = ""
            
            if section_id == "persoonsgegevens":
                section_specific_instructions = """
                Doorzoek de documenten zeer grondig naar de volgende persoonsgegevens:
                - Naam en geslacht van de cliënt (zoek naar dhr., mevrouw, voornaam, achternaam, heer, mw., etc.)
                - Geboortedatum (zoek naar geb., geboren, geboortedatum in diverse formaten)
                - Adresgegevens (zoek naar adres, woonplaats, postcode, straat, huisnummer)
                - Telefoonnummer en e-mailadres (zoek naar tel, telefoon, @, mail, e-mail)
                - Burgerlijke staat (zoek naar gehuwd, ongehuwd, samenwonend, etc.)
                - BSN (zoek naar burgerservicenummer, BSN, sofinummer)
                
                Zelfs als de informatie niet duidelijk als persoonsgegevens is gelabeld, zoek deze in de hele tekst. Bijvoorbeeld, een naam kan voorkomen in een zin als "In gesprek met dhr. Jansen..." of in de bestandsnaam of in de introductie. Bekijk de hele documentinhoud en zoek grondig naar deze gegevens. Zoek ook in documentnamen zoals die in de tekst worden genoemd (bijvoorbeeld "AD rapportage de heer Blonk" bevat al een naam).
                
                Format de persoonsgegevens in een overzichtelijke tabel of lijst. Als je bepaalde informatie niet kunt vinden, vermeld dit expliciet als "Niet aangetroffen in de aangeleverde documenten".
                """
            elif section_id == "werkgever_functie":
                section_specific_instructions = """
                Doorzoek de documenten zeer grondig naar de volgende informatie over werkgever en functie:
                - Naam en locatie van de huidige of laatste werkgever (zoek naar "werkgever", "werkt bij", "in dienst bij", bedrijfsnamen)
                - Functienaam en -omschrijving (zoek naar "functie", "werkzaamheden", "taak", "taken", "verantwoordelijkheden")
                - Datum in dienst (zoek naar "in dienst sinds", "aanvang", "gestart op", "werkzaam sinds")
                - Contractvorm (zoek naar "vast", "tijdelijk", "fulltime", "parttime", "fte", "dienstverband")
                - Werktijden/uren per week (zoek naar "uren", "werktijden", "werkdagen", "fte")
                - Status dienstverband (zoek naar "ziek", "verzuim", "actief", "werkzaam", "arbeidsongeschikt")
                - Salarisindicatie (zoek naar "salaris", "inkomen", "loon", "verdiende", "verdient")
                
                Zoek ook naar verwijzingen naar werkgever en functie in context, zoals "werkt als [functie] bij [bedrijf]" of "is werkzaam bij [bedrijf] in de functie van [functie]". Bekijk ook de bestandsnamen die kunnen informatie bevatten (bijvoorbeeld "AD rapportage mevrouw Jansen Provincie Gelderland" suggereert dat Provincie Gelderland de werkgever is).
                
                Geef een zakelijke en beknopte beschrijving van de werkgever en functie. Als je bepaalde informatie niet kunt vinden, geef dit duidelijk aan als "Niet aangetroffen in de aangeleverde documenten".
                """
            elif section_id == "aanleiding":
                section_specific_instructions = """
                Beschrijf duidelijk de aanleiding voor het arbeidsdeskundig onderzoek. Denk hierbij aan:
                - Verzuim of arbeidsongeschiktheid
                - Verzoek van werkgever, werknemer of verzekeraar
                - Wettelijk kader (Wet verbetering poortwachter, WIA, etc.)
                - Doelstelling van het onderzoek
                
                Wees specifiek over de aanleiding en het doel van het onderzoek.
                """
            elif section_id == "arbeidsverleden":
                section_specific_instructions = """
                Doorzoek de documenten nauwkeurig naar informatie over het arbeidsverleden en de opleidingsachtergrond. Zoek naar:
                
                Voor opleidingen (zoek naar woorden als "opleiding", "diploma", "studie", "school", "MBO", "HBO", "universitair"):
                - Hoogst behaalde opleiding en niveau (MBO, HBO, WO, etc.)
                - Studierichting/vakgebied
                - Jaar van afstuderen of periode van studie
                - Aanvullende cursussen, certificaten of diploma's
                
                Voor werkervaring (zoek naar "werkervaring", "carrière", "loopbaan", "functies", "gewerkt bij", "werkzaam geweest"):
                - Chronologisch overzicht van eerdere dienstverbanden
                - Periode (van - tot) per functie
                - Werkgevers
                - Functietitels
                - Verantwoordelijkheden en taken
                
                Voor vaardigheden en competenties (zoek naar "vaardigheden", "competenties", "kwaliteiten", "sterktes", "capaciteiten"):
                - Vaktechnische vaardigheden
                - Sociale vaardigheden
                - Computervaardigheden
                - Talen
                
                Besteed ook aandacht aan de context waarin deze informatie kan voorkomen, zoals in een CV-achtig overzicht of in lopende tekst. Zoek deze informatie in het gehele document, inclusief bijlagen of CV's die mogelijk zijn ingevoegd.
                
                Geef een chronologisch overzicht van opleidingen (van hoogst behaalde naar lager/ouder) en werkervaring (van meest recent naar ouder). Vermeld bij opleidingen het niveau en de discipline, bij werkervaring de periode, werkgever, functie en verantwoordelijkheden. Als bepaalde informatie niet te vinden is, geef dit dan duidelijk aan.
                """
            elif section_id == "medische_situatie":
                section_specific_instructions = """
                Beschrijf objectief de medische situatie van cliënt op basis van de documenten. Beperk je tot:
                - Diagnose en aandoeningen (alleen indien vermeld in documenten)
                - Behandeltraject en zorgverleners
                - Prognose (indien vermeld door medisch specialist/bedrijfsarts)
                - Impact op dagelijks functioneren
                
                Vermijd medische interpretaties of aannames die niet in de documenten staan.
                """
            elif section_id == "belastbaarheid":
                section_specific_instructions = """
                Analyseer de belastbaarheid van de cliënt op basis van medische informatie en functionele mogelijkheden. Bespreek:
                - Fysieke belastbaarheid (staan, zitten, tillen, etc.)
                - Mentale belastbaarheid (concentratie, stressbestendigheid, etc.)
                - Sociale belastbaarheid (samenwerken, klantcontact, etc.)
                - Mogelijke aanpassingen om belastbaarheid te vergroten
                
                Baseer je analyse op concrete informatie uit de documenten, zoals FML of inzetbaarheidsprofiel indien beschikbaar.
                """
            elif section_id == "belasting_huidige_functie":
                section_specific_instructions = """
                Beschrijf de belasting in de huidige/laatst uitgeoefende functie en vergelijk deze met de belastbaarheid:
                - Fysieke belasting (werkhouding, bewegingen, krachtsuitoefening)
                - Mentale belasting (concentratie, deadlines, verantwoordelijkheden)
                - Sociale belasting (teamwork, klantcontact, leidinggeven)
                - Omgevingsfactoren (geluid, klimaat, werkplek)
                
                Concludeer of er sprake is van disbalans tussen belasting en belastbaarheid en waar deze zich bevindt.
                """
            elif section_id == "visie_ad":
                section_specific_instructions = """
                Geef een professionele visie op de arbeidsmogelijkheden, rekening houdend met:
                - Balans tussen belasting en belastbaarheid
                - Mogelijkheden voor werkhervatting of aanpassingen
                - Beperkingen en oplossingsrichtingen
                - Benodigde interventies of ondersteuning
                
                Wees concreet en onderbouw je professionele mening met feiten uit het dossier.
                """
            elif section_id == "matching":
                section_specific_instructions = """
                Beschrijf de matching overwegingen met betrekking tot passend werk:
                - Welke functies/werkzaamheden zijn passend gezien de belastbaarheid?
                - Welke aanpassingen zijn nodig in huidige functie?
                - Alternatieven binnen het bedrijf (indien van toepassing)
                - Kansrijke beroepen of sectoren (indien re-integratie 2e spoor relevant is)
                
                Wees specifiek en realistisch in je advies over passend werk.
                """
            elif section_id == "conclusie":
                section_specific_instructions = """
                Formuleer een duidelijke conclusie en concreet advies:
                - Korte samenvatting van de belangrijkste bevindingen
                - Concrete adviezen voor werkgever en werknemer
                - Beantwoording van de onderzoeksvraag
                - Advies over re-integratiemogelijkheden (1e of 2e spoor)
                - Vervolgstappen en tijdpad
                
                Formuleer het advies puntsgewijs en maak concrete, actionable aanbevelingen.
                """
            elif section_id == "samenvatting":
                section_specific_instructions = """
                Maak een bondige samenvatting van het gehele rapport met de belangrijkste elementen:
                - Aanleiding en doel van het onderzoek
                - Kernproblematiek en beperkingen
                - Belangrijkste bevindingen t.a.v. belasting en belastbaarheid
                - Conclusie over arbeidsmogelijkheden
                - Essentie van het advies
                
                Houd de samenvatting beknopt maar volledig, maximaal één A4.
                """

            # Combineer basis prompt met specifieke instructies
            prompt = f"""{base_prompt}
            
            {section_specific_instructions}
            
            DOCUMENTEN:
            {all_content}
            """
            
            # Gemini doesn't support system role in this format
            # Add system instruction to the user prompt instead
            combined_prompt = f"{system_instruction}\n\n{prompt}"
            response = model.generate_content(combined_prompt)
            
            if response.text:
                content = response.text
            else:
                content = "Op basis van de beschikbare documenten is deze sectie samengesteld. Voor meer specifieke informatie zijn aanvullende documenten gewenst."
                
            # Prepare section result
            section_result = {
                "content": content,
                "chunk_ids": [],  # No specific chunks used
                "prompt": prompt
            }
            
        except Exception as direct_error:
            print(f"Direct approach failed: {str(direct_error)}")
            
            # Fallback to RAG pipeline
            try:
                from app.tasks.generate_report_tasks.rag_pipeline import generate_content_for_section
                
                # Generate content using RAG pipeline
                section_result = generate_content_for_section(
                    section_id=section_id,
                    section_info=section_info,
                    document_ids=document_ids,
                    case_id=case_id
                )
            except Exception as rag_error:
                # If both approaches fail, return a simple error message as content
                section_result = {
                    "content": "Op basis van de beschikbare documenten is een objectieve analyse gemaakt. Voor meer specifieke informatie zijn aanvullende documenten gewenst.",
                    "chunk_ids": [],
                    "prompt": "Beide benaderingen (direct en RAG) zijn mislukt."
                }
        
        # Get case information and document IDs
        case_id = report["case_id"]
        
        # Get processed documents for the case
        document_response = db_service.get_documents_for_case(case_id)
        document_ids = [doc["id"] for doc in document_response if doc["status"] == "processed"]
        
        if not document_ids:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No processed documents found for this case"
            )
        
        # Get section info from template
        section_id = section_data.section_id
        section_info = template["sections"][section_id]
        
        # Update report status
        db_service.update_report(str(section_data.report_id), {
            "status": "generating",
        })
        
        # Use the already generated section_result for immediate update
        try:
            # Get current report content
            report_data = db_service.get_row_by_id("report", str(section_data.report_id))
            if not report_data:
                raise ValueError(f"Report {section_data.report_id} not found")
            
            # Parse content as dictionary or initialize if empty
            content = report_data.get("content", {})
            if isinstance(content, str) and content.strip() in ["", "{}"]:
                content = {}
            
            # Update the specified section
            content[section_id] = section_result["content"]
            
            # Get or initialize metadata
            metadata = report_data.get("metadata", report_data.get("report_metadata", {}))
            if isinstance(metadata, str) and metadata.strip() in ["", "{}"]:
                metadata = {}
            
            # Initialize sections metadata if not exists
            if "sections" not in metadata:
                metadata["sections"] = {}
                
            # Update the section metadata
            metadata["sections"][section_id] = {
                "regenerated_at": datetime.utcnow().isoformat(),
                "chunk_ids": section_result.get("chunk_ids", []),
                "prompt": section_result.get("prompt", ""),
                "approach": "direct_hybrid"
            }
            
            # Update report with regenerated content
            db_service.update_report(str(section_data.report_id), {
                "status": "generated",
                "content": content,
                "report_metadata": metadata
            })
            
            return {
                "status": "success",
                "message": f"Successfully regenerated section '{section_id}'",
                "report_id": str(section_data.report_id),
                "section_id": section_id
            }
                
        except Exception as e:
            # Update report to mark error
            db_service.update_report(str(section_data.report_id), {
                "status": "generated",  # Keep as generated since only one section failed
                "error": f"Error updating report with regenerated section: {str(e)}"
            })
            
            # Return error response
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error updating report with regenerated section: {str(e)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error regenerating section: {str(e)}"
        )

@router.post("/{report_id}/generate-structured", response_model=Dict[str, Any])
async def generate_structured_report(
    report_id: UUID, 
    user_info = Depends(verify_token)
):
    """
    Generate a structured AD report using the new structured approach
    This will generate all data needed for the form-based template
    """
    user_id = user_info["user_id"]
    
    try:
        # Get the report
        report = db_service.get_report(str(report_id), user_id)
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Import the AD report task
        from app.tasks.generate_report_tasks.ad_report_task import generate_ad_report_task
        
        # Start the structured generation task
        task = generate_ad_report_task.apply_async(
            args=[str(report_id), True],  # use_structured_approach=True
            queue='reports'
        )
        
        # Update report status
        db_service.update_report(str(report_id), {
            "status": "generating",
            "task_id": task.id
        })
        
        return {
            "status": "generating",
            "task_id": task.id,
            "message": "Structured AD report generation started"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error starting structured report generation: {str(e)}"
        )

@router.get("/{report_id}/export/docx", response_class=FileResponse)
async def export_report_docx(
    report_id: UUID,
    layout: str = "standaard",
    token: str = None,
    user_info = Depends(verify_token)
):
    """
    Export a report to DOCX format
    """
    # We have two authentication methods:
    # 1. Regular token authentication from verify_token dependency
    # 2. Token passed as a query parameter (for direct downloads)

    # If token is provided in query and verify_token failed (returns None), use the query token
    if not user_info and token:
        # This is a simplified mock version for development
        print("Using mock user ID for development")
        user_id = "example_user_id"  # Mock user ID for development
    else:
        user_id = user_info["user_id"]

    print(f"Exporting DOCX report for user_id: {user_id}, report_id: {report_id}, layout: {layout}")

    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(report_id), user_id)

        if not report:
            print(f"Report not found: {report_id} for user: {user_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )

        print(f"Found report: {report.get('title')}, status: {report.get('status')}")

        # Create temporary directory if it doesn't exist
        temp_dir = os.path.join(settings.STORAGE_PATH, "temp")
        os.makedirs(temp_dir, exist_ok=True)

        # Generate a filename based on report title and ID
        safe_title = "".join(c for c in report["title"] if c.isalnum() or c in [' ', '_']).rstrip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}_rapport_{str(report_id)[:8]}.docx"
        output_path = os.path.join(temp_dir, filename)

        # Get user profile for the report
        print(f"Fetching user profile for user_id: {user_id}")
        user_profile = db_service.get_user_profile(user_id)
        print(f"User profile found: {user_profile is not None}")

        if user_profile:
            print(f"Profile details: {user_profile.get('first_name')} {user_profile.get('last_name')}, {user_profile.get('company_name')}")
            print(f"Complete user profile: {user_profile}")
        else:
            print("WARNING: No user profile found, this will affect document formatting")

        # Add user profile and layout to report metadata
        report_metadata = report.get("metadata", report.get("report_metadata", {}))
        if not report_metadata:
            report_metadata = {}

        # ALWAYS add user profile to metadata for consistent behavior
        if user_profile:
            print("Adding user profile to report metadata")
            report_metadata["user_profile"] = user_profile
            # Also add directly to report for guaranteed access
            report["user_profile"] = user_profile
        else:
            print("WARNING: No user profile available to add to metadata")

        # Override template type with the requested layout
        print(f"Setting template type to: {layout}")
        report_metadata["template_type"] = layout

        # Update report metadata in all possible locations for maximum compatibility
        report["metadata"] = report_metadata
        report["report_metadata"] = report_metadata

        # Print the final structure to verify
        print(f"Final report structure keys: {list(report.keys())}")
        print(f"Final metadata structure keys: {list(report_metadata.keys()) if report_metadata else 'None'}")

        # Generate the DOCX file
        print(f"Generating DOCX file at: {output_path}")
        try:
            # Try enhanced export first, fallback to original if needed
            try:
                from app.utils.enhanced_document_export import export_report_to_docx_enhanced
                file_path = export_report_to_docx_enhanced(report, output_path, layout)
                print("Using enhanced DOCX export")
            except ImportError:
                print("Enhanced export not available, using original export")
                from app.utils.document_export import export_report_to_docx
                file_path = export_report_to_docx(report, output_path)

            # Verify the file exists
            if not os.path.exists(file_path):
                print(f"DOCX file was not created at expected path: {file_path}")
                raise FileNotFoundError(f"Generated file not found at {file_path}")

            file_size = os.path.getsize(file_path)
            print(f"DOCX file generated successfully at: {file_path}, size: {file_size} bytes")

            if file_size == 0:
                print("WARNING: Generated file has zero size")
                raise ValueError("Generated file has zero size")

            # Return the file as a download
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as doc_error:
            print(f"Error generating DOCX file: {str(doc_error)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")

            # Create a simple error DOCX as fallback
            try:
                print("Attempting to create a simple fallback DOCX file")
                # Create a simple document with error information
                from docx import Document
                doc = Document()
                doc.add_heading('Rapport generatie fout', 0)
                doc.add_paragraph('Er is een fout opgetreden bij het genereren van het rapport.')
                doc.add_paragraph(f'Rapport titel: {report.get("title", "Onbekend")}')
                doc.add_paragraph(f'Rapport ID: {report_id}')
                doc.add_paragraph(f'Datum: {datetime.now().strftime("%d-%m-%Y %H:%M:%S")}')
                doc.add_paragraph('Probeer het later opnieuw of neem contact op met de beheerder.')

                # Save simple document
                fallback_path = os.path.join(temp_dir, f"error_{filename}")
                doc.save(fallback_path)
                print(f"Fallback DOCX created at: {fallback_path}")

                return FileResponse(
                    path=fallback_path,
                    filename=f"fout_{filename}",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as fallback_error:
                print(f"Even fallback document creation failed: {str(fallback_error)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Error exporting report: {str(doc_error)}. Fallback also failed: {str(fallback_error)}"
                )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Error exporting report: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting report: {str(e)}"
        )

@router.get("/{report_id}/export/pdf", response_class=FileResponse)
async def export_report_pdf(
    report_id: UUID,
    layout: str = "standaard",
    token: str = None,
    user_info = Depends(verify_token)
):
    """
    Export a report to PDF format
    """
    # Authentication handling (same as DOCX export)
    if not user_info and token:
        print("Using mock user ID for development")
        user_id = "example_user_id"
    else:
        user_id = user_info["user_id"]

    print(f"Exporting PDF report for user_id: {user_id}, report_id: {report_id}, layout: {layout}")

    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(report_id), user_id)

        if not report:
            print(f"Report not found: {report_id} for user: {user_id}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )

        print(f"Found report: {report.get('title')}, status: {report.get('status')}")

        # Create temporary directory if it doesn't exist
        temp_dir = os.path.join(settings.STORAGE_PATH, "temp")
        os.makedirs(temp_dir, exist_ok=True)

        # Generate a filename based on report title and ID
        safe_title = "".join(c for c in report["title"] if c.isalnum() or c in [' ', '_']).rstrip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}_rapport_{str(report_id)[:8]}.pdf"
        output_path = os.path.join(temp_dir, filename)

        # Get user profile for the report
        print(f"Fetching user profile for user_id: {user_id}")
        user_profile = db_service.get_user_profile(user_id)
        print(f"User profile found: {user_profile is not None}")

        if user_profile:
            print(f"Profile details: {user_profile.get('first_name')} {user_profile.get('last_name')}, {user_profile.get('company_name')}")
        else:
            print("WARNING: No user profile found, this will affect document formatting")

        # Add user profile and layout to report metadata
        report_metadata = report.get("metadata", report.get("report_metadata", {}))
        if not report_metadata:
            report_metadata = {}

        # Add user profile to metadata for consistent behavior
        if user_profile:
            print("Adding user profile to report metadata")
            report_metadata["user_profile"] = user_profile
            report["user_profile"] = user_profile
        else:
            print("WARNING: No user profile available to add to metadata")

        # Set template type
        print(f"Setting template type to: {layout}")
        report_metadata["template_type"] = layout

        # Update report metadata
        report["metadata"] = report_metadata
        report["report_metadata"] = report_metadata

        print(f"Final report structure keys: {list(report.keys())}")
        print(f"Final metadata structure keys: {list(report_metadata.keys()) if report_metadata else 'None'}")

        # Generate the PDF file
        print(f"Generating PDF file at: {output_path}")
        try:
            from app.utils.enhanced_document_export import export_report_to_pdf
            file_path = export_report_to_pdf(report, output_path, layout)

            # Verify the file exists
            if not os.path.exists(file_path):
                print(f"PDF file was not created at expected path: {file_path}")
                raise FileNotFoundError(f"Generated file not found at {file_path}")

            file_size = os.path.getsize(file_path)
            print(f"PDF file generated successfully at: {file_path}, size: {file_size} bytes")

            if file_size == 0:
                print("WARNING: Generated file has zero size")
                raise ValueError("Generated file has zero size")

            # Return the file as a download
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="application/pdf"
            )
        except ImportError as ie:
            print(f"PDF export not available: {str(ie)}")
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail="PDF export requires additional packages. Please install reportlab."
            )
        except Exception as doc_error:
            print(f"Error generating PDF file: {str(doc_error)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")

            # Fallback: create error message
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error generating PDF: {str(doc_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"Error exporting PDF report: {str(e)}")
        print(f"Error type: {type(e).__name__}")
        print(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting PDF report: {str(e)}"
        )

@router.get("/{report_id}/export/html", response_class=FileResponse)
async def export_report_html(
    report_id: UUID,
    layout: str = "standaard",
    token: str = None,
    user_info = Depends(verify_token)
):
    """
    Export a report to HTML format for web viewing
    """
    # Authentication handling (same as other exports)
    if not user_info and token:
        print("Using mock user ID for development")
        user_id = "example_user_id"
    else:
        user_id = user_info["user_id"]

    print(f"Exporting HTML report for user_id: {user_id}, report_id: {report_id}, layout: {layout}")

    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(report_id), user_id)

        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )

        # Create temporary directory
        temp_dir = os.path.join(settings.STORAGE_PATH, "temp")
        os.makedirs(temp_dir, exist_ok=True)

        # Generate filename
        safe_title = "".join(c for c in report["title"] if c.isalnum() or c in [' ', '_']).rstrip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}_rapport_{str(report_id)[:8]}.html"
        output_path = os.path.join(temp_dir, filename)

        # Get user profile
        user_profile = db_service.get_user_profile(user_id)

        # Prepare report data
        report_metadata = report.get("metadata", report.get("report_metadata", {}))
        if not report_metadata:
            report_metadata = {}

        if user_profile:
            report_metadata["user_profile"] = user_profile
            report["user_profile"] = user_profile

        report_metadata["template_type"] = layout
        report["metadata"] = report_metadata
        report["report_metadata"] = report_metadata

        # Generate HTML file
        try:
            from app.utils.enhanced_document_export import export_report_to_html
            file_path = export_report_to_html(report, output_path, layout)

            # Verify file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Generated file not found at {file_path}")

            # Return the file
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="text/html"
            )
        except ImportError as ie:
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail=f"HTML export requires additional packages: {str(ie)}"
            )
        except Exception as doc_error:
            print(f"Error generating HTML file: {str(doc_error)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error generating HTML: {str(doc_error)}"
            )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error exporting HTML report: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting HTML report: {str(e)}"
        )

@router.get("/export/templates", response_model=Dict[str, Any])
async def get_export_templates():
    """
    Get available export templates with descriptions
    """
    return {
        "templates": {
            "standaard": {
                "id": "standaard",
                "name": "Standaard",
                "description": "Traditional Dutch business format with professional black section headers",
                "features": [
                    "Classic layout with centered title page",
                    "Professional typography with Times New Roman",
                    "Standard margins and spacing",
                    "Essential contact information"
                ]
            },
            "modern": {
                "id": "modern", 
                "name": "Modern",
                "description": "Clean contemporary layout with blue header hierarchy and headers/footers",
                "features": [
                    "Two-column title page layout",
                    "Professional headers and footers",
                    "Comprehensive profile information",
                    "Modern color scheme with blue accents"
                ]
            },
            "professioneel": {
                "id": "professioneel",
                "name": "Professioneel", 
                "description": "Minimalist professional design with structured blue subsection headers",
                "features": [
                    "Minimalist and clean design",
                    "Professional separator elements",
                    "Condensed profile information",
                    "Focus on content readability"
                ]
            },
            "compact": {
                "id": "compact",
                "name": "Compact",
                "description": "Space-efficient format with smaller blue headers for concise reports",
                "features": [
                    "Smaller font sizes for efficiency",
                    "Reduced margins for more content",
                    "Single-line profile information",
                    "Optimized for shorter reports"
                ]
            }
        },
        "formats": {
            "docx": {
                "name": "Microsoft Word",
                "description": "Professional Word document with full formatting",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "extension": ".docx"
            },
            "pdf": {
                "name": "PDF Document", 
                "description": "Portable document format for universal viewing",
                "mime_type": "application/pdf",
                "extension": ".pdf"
            },
            "html": {
                "name": "HTML Web Page",
                "description": "Web page format for online viewing or further conversion",
                "mime_type": "text/html", 
                "extension": ".html"
            }
        }
    }

@router.post("/debug/create", status_code=status.HTTP_201_CREATED)
async def debug_create_report(user_info = Depends(verify_token)):
    """
    Debug endpoint to create a test report with minimal parameters
    """
    user_id = user_info["user_id"]

    try:
        # Create a simple test report
        import uuid
        case_id = "c452169e-d5ab-4807-8985-b76cb51a2b7e"  # Use your actual test case ID
        title = f"Test Report {uuid.uuid4().hex[:8]}"
        template_id = "staatvandienst"

        # Debug: direct call to db_service
        print("Calling db_service.create_report directly...")
        report = db_service.create_report(
            case_id=case_id,
            user_id=user_id,
            title=title,
            template_id=template_id,
            layout_type="standaard"
        )

        if not report:
            return {"status": "error", "message": "Failed to create report directly"}

        return {"status": "success", "report": report}

    except Exception as e:
        import traceback
        print(f"Debug create report error: {str(e)}")
        print(traceback.format_exc())
        return {"status": "error", "message": str(e)}

@router.delete("/{report_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_report(report_id: UUID, user_info = Depends(verify_token)):
    """
    Delete a report
    """
    user_id = user_info["user_id"]
    
    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
            
        # Delete report record
        success = db_service.delete_row("report", str(report_id))
        
        if not success:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to delete report"
            )
            
        return None
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting report: {str(e)}"
        )

# NEW STRUCTURED OUTPUT ENDPOINTS

@router.post("/generate-structured-section", response_model=Dict[str, Any])
async def generate_structured_section(
    section_data: ReportSectionGenerate,
    output_format: str = "structured",  # structured, html, markdown, plain, json
    user_info = Depends(verify_token)
):
    """
    Generate a report section with structured output
    
    Supported output formats:
    - structured: Full structured data with metadata
    - html: HTML formatted output
    - markdown: Markdown formatted output  
    - plain: Plain text (backwards compatible)
    - json: Raw JSON data
    """
    user_id = user_info["user_id"]
    
    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(section_data.report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Check if section exists in the template
        template_id = report["template_id"]
        if template_id not in MVP_TEMPLATES:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Template not found"
            )
            
        template = MVP_TEMPLATES[template_id]
        if section_data.section_id not in template["sections"]:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Section '{section_data.section_id}' not found in template"
            )
        
        # Get case information and document IDs
        case_id = report["case_id"]
        document_response = db_service.get_documents_for_case(case_id)
        document_ids = [doc["id"] for doc in document_response if doc["status"] == "processed"]
        
        if not document_ids:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No processed documents found for this case"
            )
        
        # Get user profile for context
        try:
            user_profile = db_service.get_row("user_profile", {"user_id": user_id})
        except Exception:
            user_profile = None
        
        # Get section info from template
        section_info = template["sections"][section_data.section_id]
        
        # Generate structured content using the new pipeline
        result = await generate_structured_content_for_section(
            section_id=section_data.section_id,
            section_info=section_info,
            document_ids=document_ids,
            case_id=case_id,
            user_profile=user_profile,
            output_format=output_format
        )
        
        return {
            "success": True,
            "section_id": section_data.section_id,
            "output_format": output_format,
            "content": result["content"],
            "structured_content": result.get("structured_content"),
            "metadata": result.get("metadata", {}),
            "chunk_ids": result.get("chunk_ids", []),
            "error": result.get("error")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating structured section: {str(e)}"
        )

@router.get("/{report_id}/structured", response_model=Dict[str, Any])
async def get_report_structured(
    report_id: UUID, 
    output_format: str = "structured",
    user_info = Depends(verify_token)
):
    """
    Get a report with structured output format
    
    Returns the report content in the requested structured format
    """
    user_id = user_info["user_id"]
    
    try:
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Check if report has structured content metadata
        metadata = report.get("metadata", {})
        structured_sections = {}
        
        # If report content has structured metadata, use it
        if "structured_content" in metadata:
            structured_sections = metadata["structured_content"]
        
        # Standardize report structure to professional format
        content = report.get("content", {})
        if isinstance(content, dict):
            # Apply professional section ordering and mapping
            content = standardize_report_structure(content, "staatvandienst")
            
            from app.utils.structured_output_generator import OutputFormatter
            formatter = OutputFormatter()
            
            # Get professional section order
            ordered_section_ids = get_ordered_sections("staatvandienst")
            converted_sections = {}
            
            # Process sections in professional order
            for section_id in ordered_section_ids:
                section_content = content.get(section_id, "")
                
                if section_id in structured_sections:
                    # Use existing structured data
                    structured_data = structured_sections[section_id]
                    if output_format == "html":
                        converted_sections[section_id] = formatter.to_html(structured_data)
                    elif output_format == "markdown":
                        converted_sections[section_id] = formatter.to_markdown(structured_data)
                    elif output_format == "json":
                        converted_sections[section_id] = formatter.to_json(structured_data)
                    else:
                        converted_sections[section_id] = structured_data
                elif section_content:
                    # Use plain text content
                    converted_sections[section_id] = section_content
            
            content = converted_sections
        
        return {
            "id": report["id"],
            "title": report.get("title", ""),
            "status": report.get("status", ""),
            "output_format": output_format,
            "content": content,
            "structured_metadata": structured_sections,
            "template_id": report.get("template_id", ""),
            "created_at": report.get("created_at", ""),
            "updated_at": report.get("updated_at", "")
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error retrieving structured report: {str(e)}"
        )

@router.post("/{report_id}/convert-format", response_model=Dict[str, Any])
async def convert_report_format(
    report_id: UUID,
    target_format: str,  # html, markdown, json, plain
    user_info = Depends(verify_token)
):
    """
    Convert existing report content to different structured formats
    """
    user_id = user_info["user_id"]
    
    try:
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Check if we have structured content to convert
        metadata = report.get("metadata", {})
        if "structured_content" not in metadata:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Report does not have structured content available for conversion"
            )
        
        from app.utils.structured_output_generator import OutputFormatter, SectionContent
        formatter = OutputFormatter()
        
        structured_sections = metadata["structured_content"]
        converted_sections = {}
        
        for section_id, structured_data in structured_sections.items():
            try:
                # Reconstruct SectionContent object
                section_content = SectionContent(**structured_data)
                
                # Convert to target format
                if target_format == "html":
                    converted_sections[section_id] = formatter.to_html(section_content)
                elif target_format == "markdown":
                    converted_sections[section_id] = formatter.to_markdown(section_content)
                elif target_format == "json":
                    converted_sections[section_id] = formatter.to_json(section_content)
                elif target_format == "plain":
                    converted_sections[section_id] = formatter.to_plain_text(section_content)
                else:
                    raise ValueError(f"Unsupported target format: {target_format}")
                    
            except Exception as e:
                # Fallback to original content if conversion fails
                converted_sections[section_id] = report.get("content", {}).get(section_id, "")
        
        return {
            "success": True,
            "report_id": str(report_id),
            "target_format": target_format,
            "converted_content": converted_sections,
            "conversion_timestamp": datetime.utcnow().isoformat()
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error converting report format: {str(e)}"
        )

@router.get("/{report_id}/export/{format}")
async def export_ad_report(
    report_id: UUID, 
    format: str,
    template: str = "standaard",
    user_info = Depends(verify_token)
):
    """
    Export AD report in specified format using new structured system
    
    Args:
        report_id: UUID of the report to export
        format: Export format (markdown, html, pdf, json)
        template: Template style (standaard, modern, professioneel, compact)
        
    Returns:
        Report in requested format or file response
    """
    user_id = user_info["user_id"]
    
    try:
        # Get report from database
        report = db_service.get_report(str(report_id), user_id)
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Check ownership
        if report.get("user_id") != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        # Check if report has structured data
        content = report.get("content", {})
        if not content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Report has no content"
            )
        
        # Try to get structured data first
        structured_data = content.get("structured_data")
        
        if structured_data:
            # Use new structured approach
            from app.models.ad_report_structure import ADReport
            from app.utils.ad_report_renderer import ADReportRenderer
            
            try:
                # Reconstruct AD report from structured data
                ad_report = ADReport(**structured_data)
                
                # Render in requested format
                renderer = ADReportRenderer(template=template)
                
                if format.lower() == "markdown":
                    rendered_content = renderer.render_markdown(ad_report)
                    return {"content": rendered_content, "format": "markdown"}
                    
                elif format.lower() == "html":
                    rendered_content = renderer.render_html(ad_report)
                    return {"content": rendered_content, "format": "html"}
                    
                elif format.lower() == "json":
                    rendered_content = renderer.render_json(ad_report)
                    return {"content": rendered_content, "format": "json"}
                    
                else:
                    raise HTTPException(
                        status_code=status.HTTP_400_BAD_REQUEST,
                        detail=f"Unsupported format: {format}. Supported: markdown, html, json"
                    )
                    
            except Exception as e:
                logger.error(f"Error rendering structured report: {str(e)}")
                # Fall back to legacy sections approach
                
        # Fallback to legacy sections if no structured data
        sections = content.get("sections", {})
        if not sections:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Report has no sections or structured data"
            )
        
        # Convert legacy sections using adapter
        from app.utils.ad_pipeline_adapter import ADPipelineAdapter
        
        adapter = ADPipelineAdapter()
        case = db_service.get_case(str(report.get("case_id")), user_id)
        
        case_data = {
            "client_name": case.get("client_name", ""),
            "company_name": case.get("company_name", "")
        } if case else {}
        
        # Convert to structured format
        import asyncio
        ad_report = asyncio.run(
            adapter.generate_ad_report_from_legacy_sections(
                context="",  # No context needed for conversion
                sections=sections,
                case_data=case_data,
                template_id=template
            )
        )
        
        # Render
        rendered_content = await adapter.render_ad_report(ad_report, format, template)
        
        return {
            "content": rendered_content,
            "format": format,
            "template": template,
            "report_id": str(report_id)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        logger.error(f"Error exporting AD report: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting report: {str(e)}"
        )

@router.post("/{report_id}/regenerate-section")
async def regenerate_report_section(
    report_id: UUID,
    section_data: ReportSectionGenerate,
    use_structured: bool = True,
    user_info = Depends(verify_token)
):
    """
    Regenerate a specific section of an AD report
    
    Args:
        report_id: UUID of the report
        section_data: Section regeneration request
        use_structured: Whether to use structured AD approach
        
    Returns:
        Task information for section regeneration
    """
    user_id = user_info["user_id"]
    
    try:
        # Get report and verify ownership
        report = db_service.get_report(str(report_id), user_id)
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        if report.get("user_id") != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        # Start section regeneration task
        task = celery.send_task(
            "generate_ad_report_section",
            args=[str(report_id), section_data.section_id, use_structured]
        )
        
        return {
            "success": True,
            "task_id": task.id,
            "report_id": str(report_id),
            "section_id": section_data.section_id,
            "message": f"Started regeneration of section {section_data.section_id}"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error regenerating section: {str(e)}"
        )

@router.get("/{report_id}/structure")
async def get_report_structure(
    report_id: UUID,
    user_info = Depends(verify_token)
):
    """
    Get the structure information of an AD report
    
    Returns metadata about sections, structured data availability, etc.
    """
    user_id = user_info["user_id"]
    
    try:
        report = db_service.get_report(str(report_id), user_id)
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        if report.get("user_id") != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        content = report.get("content", {})
        metadata = report.get("metadata", {})
        
        # Check what's available
        has_structured_data = "structured_data" in content
        has_sections = "sections" in content
        has_markdown = "markdown" in content
        has_html = "html" in content
        
        structure_info = {
            "report_id": str(report_id),
            "status": report.get("status", "unknown"),
            "template_id": report.get("template_id"),
            "has_structured_data": has_structured_data,
            "has_legacy_sections": has_sections,
            "has_markdown": has_markdown,
            "has_html": has_html,
            "generation_method": metadata.get("generation_method", "unknown"),
            "sections_count": metadata.get("sections_count", 0),
            "fml_rubrieken_count": metadata.get("fml_rubrieken_count", 0),
            "available_formats": [],
            "available_sections": []
        }
        
        # Determine available formats
        if has_structured_data or has_sections:
            structure_info["available_formats"] = ["markdown", "html", "json"]
        
        # List available sections
        if has_sections:
            structure_info["available_sections"] = list(content["sections"].keys())
        elif has_structured_data:
            # Determine sections from structured data
            structured_sections = [
                "vraagstelling", "ondernomen_activiteiten", "voorgeschiedenis",
                "gegevens_werkgever", "gegevens_werknemer", "belastbaarheid", 
                "eigen_functie", "visie_arbeidsdeskundige", "conclusie"
            ]
            structure_info["available_sections"] = structured_sections
        
        return structure_info
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting report structure: {str(e)}"
        )


@router.post("/{report_id}/generate-ad-structure", response_model=Dict[str, Any])
async def generate_ad_structure(
    report_id: UUID,
    user_info = Depends(verify_token)
):
    """
    Generate a complete ADReport structure using Pydantic models
    This creates structured JSON data that perfectly matches the ADReport schema
    """
    user_id = user_info["user_id"]
    
    try:
        # Get the report
        report = db_service.get_report(str(report_id), user_id)
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Get case information
        case_id = report["case_id"]
        
        # Get all document content for this case
        # Fix: Use IN operator instead of array comparison
        try:
            documents = db_service.get_rows("document", {
                "case_id": case_id
            })
            # Filter documents in Python instead of SQL to avoid array issues
            documents = [doc for doc in documents if doc.get("status") in ["processed", "enhanced"]]
        except Exception as e:
            logger.warning(f"Error fetching documents: {e}, using empty list")
            documents = []
        
        # Collect all document text
        all_content = ""
        for doc in documents:
            chunks = db_service.get_document_chunks(doc["id"])
            for chunk in chunks:
                if chunk.get("content"):
                    all_content += chunk["content"] + "\n"
        
        # Get user profile for adviseur data
        user_profile = db_service.get_user_profile(user_id)
        
        # Generate structured AD report using the configured LLM provider
        from app.utils.llm_provider import create_llm_instance, get_llm_model_name
        
        model = create_llm_instance(
            temperature=0.1,
            max_tokens=8192,
            dangerous_content_level="BLOCK_NONE"
        )
        
        # Create ADReport template structure for the prompt
        empty_report = ADReportGenerator.create_empty_report()
        
        # Prepare the prompt for structured data generation
        system_prompt = """Je bent een professionele arbeidsdeskundige die een volledig gestructureerd AD rapport opstelt.
Je moet een JSON object maken dat PRECIES voldoet aan het ADReport Pydantic schema.

KRITIEKE REGELS - VOLG DEZE EXACT:
1. Gebruik ALLEEN de exacte veldnamen zoals hieronder gespecificeerd
2. GEEN creativiteit in de JSON structuur - volg het template PRECIES
3. Als informatie niet beschikbaar is: gebruik "[Te bepalen]" of "[Niet vermeld in documenten]"
4. Alle datums in DD-MM-YYYY formaat
5. Arrays moeten arrays blijven, objecten moeten objecten blijven
6. Geef ALLEEN een geldig JSON object terug, GEEN tekst ervoor of erna
7. ESSENTIEEL: Gebruik de EXACTE structuur uit het voorbeeld hieronder"""

        user_prompt = f"""Analyseer de documenten en maak een compleet gestructureerd AD rapport.

DOCUMENTEN INHOUD:
{all_content[:8000] if all_content else "Geen documentinhoud beschikbaar"}

ARBEIDSDESKUNDIGE PROFIEL:
{json.dumps(user_profile, indent=2, ensure_ascii=False, default=str) if user_profile else "Geen profiel beschikbaar"}

INSTRUCTIES VOOR DATAEXTRACTIE:
1. Haal werknemernaam uit functieomschrijving of medische documenten (NIET uit bedrijfsnaam)
2. Zoek specifieke bedrijfsgegevens, adressen, contactpersonen in de documenten
3. Identificeer concrete klachten, diagnoses, behandelingen uit medische rapporten
4. Vul alle "[Te bepalen]" velden in met concrete informatie uit de documenten
5. Als informatie echt niet beschikbaar is, gebruik dan "[Niet vermeld in documenten]"
6. Maak realistische arbeidsdeskundige conclusies gebaseerd op de gevonden informatie

VERPLICHTE STRUCTUUR - Genereer PRECIES dit JSON object:
{{
  "titel": "Arbeidsdeskundig rapport",
  "versie": "1.0", 
  "template": "standaard",
  "opdrachtgever": {{
    "naam_bedrijf": "[Bedrijfsnaam uit documenten]",
    "contactpersoon": "[Contactpersoon]",
    "functie_contactpersoon": "[Functie]",
    "adres": "[Adres]",
    "postcode": "[Postcode]",
    "woonplaats": "[Woonplaats]",
    "telefoonnummer": "[Telefoonnummer]",
    "email": "[Email]",
    "aard_bedrijf": "[Aard bedrijf]",
    "omvang_bedrijf": "[Omvang]"
  }},
  "werknemer": {{
    "naam": "[HAAL WERKNEMERNAAM UIT DE DOCUMENTEN - NIET DE BEDRIJFSNAAM]",
    "geboortedatum": "[Geboortedatum uit medische of intake documenten]",
    "adres": "[Woonadres uit intake formulier]",
    "postcode": "[Postcode uit intake formulier]", 
    "woonplaats": "[Woonplaats uit intake formulier]",
    "telefoonnummer": "[Telefoonnummer uit intake formulier]",
    "email": "[Email uit intake formulier]"
  }},
  "adviseur": {{
    "naam": "{user_profile.get('first_name', '') + ' ' + user_profile.get('last_name', '') if user_profile else '[Naam arbeidsdeskundige]'}",
    "functie": "{user_profile.get('job_title', 'Gecertificeerd Register Arbeidsdeskundige') if user_profile else 'Gecertificeerd Register Arbeidsdeskundige'}",
    "adres": "{user_profile.get('company_address', '[Adres organisatie]') if user_profile else '[Adres organisatie]'}",
    "postcode": "{user_profile.get('company_postal_code', '[Postcode]') if user_profile else '[Postcode]'}",
    "woonplaats": "{user_profile.get('company_city', '[Woonplaats]') if user_profile else '[Woonplaats]'}",
    "telefoonnummer": "{user_profile.get('company_phone', '[Telefoonnummer]') if user_profile else '[Telefoonnummer]'}",
    "email": "{user_profile.get('company_email', '[Email]') if user_profile else '[Email]'}"
  }},
  "onderzoek": {{
    "datum_onderzoek": "{datetime.now().strftime('%d-%m-%Y')}",
    "datum_rapportage": "{datetime.now().strftime('%d-%m-%Y')}",
    "locatie_onderzoek": "[Locatie onderzoek]"
  }},
  "samenvatting_vraagstelling": [
    "Kan werknemer het eigen werk bij de eigen werkgever nog uitvoeren?",
    "Zo nee, is het eigen werk met behulp van aanpassingen passend te maken?", 
    "Zo nee, kan werknemer ander werk bij de eigen werkgever uitvoeren?",
    "Zo nee, zijn er mogelijkheden om werknemer naar ander werk te begeleiden en is een vervolgtraject gewenst?"
  ],
  "samenvatting_conclusie": ["[Schrijf een concrete hoofdconclusie gebaseerd op de medische informatie en arbeidsgegevens uit de documenten]"],
  "vraagstelling": [
    {{"vraag": "Kan werknemer het eigen werk bij de eigen werkgever nog uitvoeren?", "antwoord": "[Geef een specifiek antwoord: Ja/Nee/Gedeeltelijk met uitleg waarom]"}},
    {{"vraag": "Zo nee, is het eigen werk met behulp van aanpassingen passend te maken?", "antwoord": "[Geef specifieke aanpassingen die mogelijk zijn]"}},
    {{"vraag": "Welke interventies zijn nodig?", "antwoord": "[Lijst concrete interventies uit de medische en functie-informatie]"}}
  ],
  "ondernomen_activiteiten": ["Dossieronderzoek medische en functiegegevens", "Analyse arbeidsomstandigheden", "Beoordeling belasting-belastbaarheid"],
  "voorgeschiedenis": "[Haal concrete medische geschiedenis, klachten, behandelingen uit de documenten - wees specifiek]",
  "verzuimhistorie": "[Haal concrete verzuimgegevens, duur, oorzaak uit de documenten]",
  "opleidingen": [],
  "arbeidsverleden_lijst": [],
  "bekwaamheden": {{"computervaardigheden": "[Te bepalen]", "taalvaardigheid": "[Te bepalen]", "rijbewijs": "[Te bepalen]"}},
  "belastbaarheid": {{
    "datum_beoordeling": "{datetime.now().strftime('%d-%m-%Y')}",
    "beoordelaar": "{user_profile.get('first_name', '') + ' ' + user_profile.get('last_name', '') if user_profile else '[Beoordelaar]'}",
    "fml_rubrieken": []
  }},
  "eigen_functie": {{
    "naam_functie": "[Huidige functie]",
    "arbeidspatroon": "[Voltijd/deeltijd]",
    "overeenkomst": "[Type contract]", 
    "aantal_uren": "[Uren per week]",
    "functieomschrijving": "[Functieomschrijving uit documenten]"
  }},
  "functiebelasting": [],
  "gesprek_werkgever": {{}},
  "gesprek_werknemer": {{}},
  "geschiktheid_eigen_werk": [
    {{
      "belastend_aspect": "[Te bepalen uit documenten]",
      "belastbaarheid_werknemer": "[Te bepalen uit documenten]", 
      "conclusie": "[Te bepalen uit documenten]"
    }}
  ],
  "conclusie_eigen_werk": "[Geef een concrete conclusie: kan wel/niet/gedeeltelijk het eigen werk doen, met specifieke reden gebaseerd op medische beperkingen vs functie-eisen]",
  "aanpassing_eigen_werk": "[Lijst specifieke werkplekaanpassingen die helpen: ergonomische hulpmiddelen, pauzeschema, aangepaste taken, etc.]",
  "geschiktheid_ander_werk_intern": "[Beoordeel of andere functies bij huidige werkgever mogelijk zijn, met voorbeelden]",
  "geschiktheid_ander_werk_extern": "[Beoordeel mogelijkheden bij andere werkgevers indien van toepassing]", 
  "visie_duurzaamheid": "[Geef prognose voor langetermijn werkhervatting: gunstig/matig/ongunstig met tijdsinschatting]",
  "trajectplan": [
    {{
      "actie": "[Concrete actie zoals 'implementeer ergonomische werkplek' of 'start gefaseerde opbouw']",
      "verantwoordelijke": "[Werkgever/Werknemer/Behandelaar/Arbeidsdeskundige]",
      "termijn": "[Realistische termijn: 1-4 weken/1-3 maanden/etc.]"
    }}
  ],
  "conclusies": [
    {{
      "conclusie": "[Hoofdconclusie over werkhervatting: mogelijk/niet mogelijk/met beperkingen]",
      "toelichting": "[Uitleg waarom deze conclusie, gebaseerd op medische informatie en functie-analyse]"
    }}
  ],
  "vervolg": ["[Concrete vervolgstappen zoals evaluatie na X weken, aanvullend onderzoek, start werkhervatting, etc.]"]
}}

ANALYTISCHE INSTRUCTIES:
1. Identificeer in de documenten: werknemernaam, functie, klachten, diagnose, behandeling
2. Maak een realistische belasting-belastbaarheid analyse
3. Concludeer specifiek of werkhervatting mogelijk is en onder welke voorwaarden
4. Geef concrete, uitvoerbare aanbevelingen
5. Gebruik Nederlandse arbeidsdeskundige terminologie

VEREISTE KWALITEIT:
- Geen placeholder tekst zoals "[Te bepalen]" tenzij informatie echt ontbreekt
- Specifieke medische diagnoses en functionele beperkingen
- Concrete werkplekaanpassingen en tijdsinschattingen
- Professionele arbeidsdeskundige taalgebruik

ESSENTIEEL: Genereer EXACT deze structuur - verander NIETS aan de veldnamen of structuur:
- opdrachtgever: OBJECT met veldnamen zoals "naam_bedrijf", "contactpersoon", etc.  
- werknemer: OBJECT met veldnamen zoals "naam", "geboortedatum", etc.
- adviseur: OBJECT met veldnamen zoals "naam", "functie", etc.
- geschiktheid_eigen_werk: ARRAY van objecten met "belastend_aspect", "belastbaarheid_werknemer", "conclusie"
- trajectplan: ARRAY van objecten met "actie", "verantwoordelijke", "termijn"  
- conclusies: ARRAY van objecten met "conclusie", "toelichting"

Genereer PRECIES de structuur hierboven - geen alternatieve interpretaties!"""
        
        # Generate content using the standardized interface
        try:
            response = model.generate_content([
                {"role": "system", "parts": [system_prompt]},
                {"role": "user", "parts": [user_prompt]}
            ])
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            error_message = str(e)
            
            # Check if it's an overloaded error
            if "529" in error_message or "overloaded" in error_message.lower():
                logger.warning("LLM API is currently overloaded, returning 503 Service Unavailable")
                raise HTTPException(
                    status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                    detail="LLM service is temporarily overloaded. Please try again in a few moments."
                )
            else:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to generate structured content: {str(e)}"
                )
        
        # Extract response text (different providers might have different response formats)
        response_text = ""
        try:
            if hasattr(response, 'text') and response.text:
                response_text = response.text
            elif hasattr(response, 'content') and response.content:
                response_text = response.content
            elif isinstance(response, str):
                response_text = response
            elif hasattr(response, 'candidates') and response.candidates:
                # Handle Google Gemini format
                response_text = response.candidates[0].content.parts[0].text
            elif hasattr(response, 'choices') and response.choices:
                # Handle OpenAI format
                response_text = response.choices[0].message.content
            else:
                logger.error(f"Unexpected response format: {type(response)}")
                logger.error(f"Response attributes: {dir(response)}")
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Failed to generate structured content - unexpected response format"
                )
        except Exception as e:
            logger.error(f"Error extracting response text: {e}")
            logger.error(f"Response type: {type(response)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to extract response text: {str(e)}"
            )
        
        if not response_text.strip():
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to generate structured content - empty response"
            )
        
        logger.info(f"LLM Response (first 500 chars): {response_text[:500]}")
        
        # Parse the JSON response - handle code blocks from Claude Sonnet 4
        try:
            # Remove markdown code blocks if present
            import re
            json_text = response_text.strip()
            
            # Extract JSON from markdown code blocks like ```json ... ```
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', json_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(1)
                logger.info("Extracted JSON from markdown code block")
            else:
                # Fallback: extract JSON object from text
                json_match = re.search(r'\{.*\}', json_text, re.DOTALL)
                if json_match:
                    json_text = json_match.group()
                    logger.info("Extracted JSON object from text")
                else:
                    # If no JSON found, check if it's a plain text response that needs fallback
                    logger.warning("No JSON structure found in response")
                    if "Op basis van de beschikbare documenten" in response_text:
                        # This is the fallback text from the hybrid processor
                        # Create a minimal structured response
                        logger.info("Creating fallback structured response")
                        structured_data = {
                            "titel": "Arbeidsdeskundig rapport",
                            "versie": "1.0",
                            "template": "standaard",
                            "opdrachtgever": {
                                "naam_bedrijf": "[Te bepalen uit documenten]",
                                "contactpersoon": "[Te bepalen]"
                            },
                            "werknemer": {
                                "naam": "[Te bepalen uit documenten]"
                            },
                            "adviseur": {
                                "naam": user_profile.get('first_name', '') + ' ' + user_profile.get('last_name', '') if user_profile else '[Naam arbeidsdeskundige]',
                                "functie": user_profile.get('job_title', 'Gecertificeerd Register Arbeidsdeskundige') if user_profile else 'Gecertificeerd Register Arbeidsdeskundige'
                            },
                            "samenvatting_conclusie": [response_text[:200] + "..."],
                            "conclusies": [{
                                "conclusie": "Analyse in uitvoering",
                                "toelichting": response_text
                            }],
                            "_fallback_response": True
                        }
                    else:
                        raise HTTPException(
                            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                            detail="No valid JSON structure found in response"
                        )
            
            if '_fallback_response' not in locals():
                structured_data = json.loads(json_text)
                
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON response: {e}")
            logger.error(f"JSON text: {json_text[:500]}...")
            
            # Create fallback structured response
            logger.info("Creating fallback structured response due to JSON parse error")
            structured_data = {
                "titel": "Arbeidsdeskundig rapport",
                "versie": "1.0",
                "template": "standaard",
                "opdrachtgever": {
                    "naam_bedrijf": "[Te bepalen uit documenten]"
                },
                "werknemer": {
                    "naam": "[Te bepalen uit documenten]"
                },
                "adviseur": {
                    "naam": user_profile.get('first_name', '') + ' ' + user_profile.get('last_name', '') if user_profile else '[Naam arbeidsdeskundige]',
                    "functie": user_profile.get('job_title', 'Gecertificeerd Register Arbeidsdeskundige') if user_profile else 'Gecertificeerd Register Arbeidsdeskundige'
                },
                "samenvatting_conclusie": ["Rapport wordt nog geanalyseerd"],
                "conclusies": [{
                    "conclusie": "Analyse in uitvoering",
                    "toelichting": "Er werd een fout gedetecteerd bij het genereren van de gestructureerde data. Het rapport kan handmatig worden aangevuld."
                }],
                "_parsing_error": str(e),
                "_fallback_response": True
            }
        
        # Validate against Pydantic model
        try:
            ad_report = ADReport(**structured_data)
        except Exception as e:
            logger.error(f"Pydantic validation failed: {e}")
            # Still return the data even if validation fails, for debugging
            structured_data["_validation_error"] = str(e)
        
        # Update the report content with structured data
        current_content = report.get("content", {})
        current_content["structured_data"] = structured_data
        
        # Update report in database
        update_data = {
            "content": current_content,
            "has_structured_data": True,
            "generation_method": "structured_ad",
            "status": "generated"
        }
        
        db_service.update_report(str(report_id), update_data)
        
        return {
            "success": True,
            "report_id": str(report_id),
            "structured_data": structured_data,
            "message": "AD rapport structuur succesvol gegenereerd"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating AD structure: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating structured AD report: {str(e)}"
        )


@router.get("/{report_id}/quality-metrics", response_model=Dict[str, Any])
async def get_report_quality_metrics(report_id: UUID, user_info = Depends(verify_token)):
    """
    Get quality metrics for an enhanced AD report
    """
    user_id = user_info["user_id"]
    
    try:
        # Get report
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Extract quality metrics
        quality_metrics = report.get("quality_metrics", {})
        validation_results = report.get("report_metadata", {}).get("validation", {})
        
        # Calculate additional metrics
        metrics = {
            "report_id": str(report_id),
            "generation_method": report.get("generation_method", "unknown"),
            "format_version": report.get("format_version", "1.0"),
            "completeness_score": quality_metrics.get("completeness_score", 0.0),
            "sections_generated": quality_metrics.get("sections_generated", 0),
            "total_sections": quality_metrics.get("total_sections", 0),
            "fml_generated": quality_metrics.get("fml_generated", False),
            "fml_rubrieken_count": report.get("fml_rubrieken_count", 0),
            "has_structured_data": report.get("has_structured_data", False),
            "missing_sections": validation_results.get("missing_sections", []),
            "validation_passed": validation_results.get("is_complete", False),
            "generation_time": None
        }
        
        # Calculate generation time if available
        if "report_metadata" in report:
            metadata = report["report_metadata"]
            if "generation_started" in metadata and "generation_completed" in metadata:
                try:
                    from datetime import datetime
                    start_time = datetime.fromisoformat(metadata["generation_started"].replace('Z', '+00:00'))
                    end_time = datetime.fromisoformat(metadata["generation_completed"].replace('Z', '+00:00'))
                    metrics["generation_time"] = int((end_time - start_time).total_seconds())
                except:
                    pass
        
        return metrics
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting quality metrics: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error getting quality metrics: {str(e)}"
        )
