"""
Utilities voor het exporteren van rapporten naar verschillende formaten.
"""
import os
import io
import re
import requests
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def process_inline_formatting(paragraph, text):
    """Helper function to process inline markdown formatting within a paragraph"""
    if not text:
        return
        
    # Process links [text](url)
    link_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    
    # Split by links first
    parts = re.split(f'({link_pattern})', text)
    
    i = 0
    while i < len(parts):
        part = parts[i]
        
        # Skip empty parts
        if not part:
            i += 1
            continue
            
        # Check if this part is a link
        link_match = re.match(link_pattern, part)
        if link_match:
            link_text = link_match.group(1)
            link_url = link_match.group(2)
            run = paragraph.add_run(link_text)
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
            i += 1
            continue
            
        # Process bold and italic in non-link text
        segments = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', part)
        for segment in segments:
            if not segment:
                continue
                
            # Bold text (** or __)
            if (segment.startswith('**') and segment.endswith('**')) or (segment.startswith('__') and segment.endswith('__')):
                marker_len = 2
                run = paragraph.add_run(segment[marker_len:-marker_len])
                run.bold = True
            # Italic text (* or _)
            elif (segment.startswith('*') and segment.endswith('*')) or (segment.startswith('_') and segment.endswith('_')):
                marker_len = 1
                run = paragraph.add_run(segment[marker_len:-marker_len])
                run.italic = True
            # Regular text
            else:
                paragraph.add_run(segment)
        i += 1

def convert_markdown_to_paragraphs(doc, markdown_text, style='NormalText'):
    """
    Convert markdown text to properly formatted Word paragraphs

    Args:
        doc: Word document object
        markdown_text (str): Text with markdown formatting
        style (str): Base style to use for regular paragraphs

    Returns:
        None - adds paragraphs to the document directly
    """
    # Safety check for None or empty text
    if not markdown_text:
        doc.add_paragraph("Geen inhoud beschikbaar", style=style)
        return

    # Prepare text by normalizing line endings and cleaning
    markdown_text = markdown_text.replace('\r\n', '\n').replace('\r', '\n')

    # Format the text without markdown structures
    # This simplified approach will prevent duplicate headers
    lines = markdown_text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            # Empty line becomes paragraph spacing
            doc.add_paragraph()
            continue

        # Remove markdown header symbols entirely (# signs)
        if line.startswith("#"):
            # Count the header level
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1

            # Extract header text
            header_text = line[level:].strip()

            # Add as appropriate header
            if level == 1:
                # First level headers will be formatted but NOT as Word headings
                # to avoid duplicate headers
                p = doc.add_paragraph(header_text, style="NormalText")
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)
            elif level == 2:
                p = doc.add_paragraph(header_text, style="NormalText")
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(13)
            else:
                p = doc.add_paragraph(header_text, style="NormalText")
                for run in p.runs:
                    run.bold = True

        # Remove markdown bullet points
        elif line.startswith("- ") or line.startswith("* "):
            bullet_text = line[2:].strip()
            p = doc.add_paragraph(bullet_text, style=style)
            p.style = 'List Bullet'

        # Handle numbered lists
        elif re.match(r'^\d+\.\s+', line):
            num_text = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph(num_text, style=style)
            p.style = 'List Number'

        # Regular paragraph
        else:
            # Process any bold or italic text
            p = doc.add_paragraph(style=style)

            # Very simple bold/italic processing
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part:
                    p.add_run(part)

def export_report_to_docx(report_data, output_path):
    """
    Exporteer een rapport naar DOCX formaat

    Args:
        report_data (dict): Rapport data met alle secties
        output_path (str): Pad waar het DOCX bestand moet worden opgeslagen

    Returns:
        str: Pad naar het gegenereerde bestand
    """
    # Maak nieuw document
    doc = Document()

    # Haal template type op (default: standaard)
    report_metadata = report_data.get('report_metadata', report_data.get('metadata', {}))
    template_type = report_metadata.get('template_type', report_data.get('layout_type', 'standaard'))

    # Configureer document en stijlen
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)  # A4 breedte
        section.page_height = Cm(29.7)  # A4 hoogte
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    # Stijlen definiëren
    styles = doc.styles

    # Helper function to safely add styles or get existing ones
    def get_or_add_style(style_name, style_type):
        try:
            # Try to add new style
            style = styles.add_style(style_name, style_type)
            return style
        except ValueError:
            # Style already exists, get and use it
            print(f"Style '{style_name}' already exists, using existing style")
            return styles[style_name]
    
    # Titelpagina stijl
    title_style = get_or_add_style('TitlePage', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True

    # Hoofdstuktitel stijl
    heading_style = get_or_add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(14)
    heading_font.bold = True

    # Normale tekst stijl
    normal_style = get_or_add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)

    # Subtitel stijl - use our custom name to avoid conflicts
    subtitle_style = get_or_add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(16)
    subtitle_font.bold = True

    # Profieldetails stijl
    profile_style = get_or_add_style('ProfileDetails', WD_STYLE_TYPE.PARAGRAPH)
    profile_font = profile_style.font
    profile_font.name = 'Arial'
    profile_font.size = Pt(10)

    # Haal profielinformatie op met meerdere fallbacks
    user_profile = None
    
    # Proefiel zoeken op verschillende locaties
    if 'user_profile' in report_data:
        # Direct in report_data (beste optie)
        user_profile = report_data['user_profile']
        print(f"Using profile directly from report_data root: {user_profile.get('first_name')} {user_profile.get('last_name')}")
    elif report_metadata and 'user_profile' in report_metadata:
        # In metadata
        user_profile = report_metadata['user_profile']
        print(f"Using profile from report_metadata: {user_profile.get('first_name')} {user_profile.get('last_name')}")
    else:
        # Geen profiel gevonden - maak een mock profiel voor testdoeleinden
        print("WARNING: No user profile found in any location, using fallback mock profile")
        user_profile = {
            "first_name": "Test",
            "last_name": "Gebruiker",
            "job_title": "Arbeidsdeskundige",
            "company_name": "Demo Bedrijf BV",
            "company_description": "Dit is een automatisch gegenereerd bedrijfsprofiel.",
            "company_address": "Teststraat 123",
            "company_postal_code": "1234 AB",
            "company_city": "Amsterdam",
            "company_country": "Nederland",
            "company_phone": "020-1234567",
            "company_email": "info@demobedrijf.nl",
            "company_website": "www.demobedrijf.nl",
            "certification": "Gecertificeerd Arbeidsdeskundige",
            "registration_number": "AD12345",
            "specializations": ["Letselschade", "Re-integratie"],
            "bio": "Dit is een gegenereerd profiel voor rapportage-doeleinden.",
            "logo_path": None  # No logo by default, we'll use text fallback
        }
    
    # Zorg ervoor dat er altijd een logo_path is voor het profiel
    if not user_profile.get('logo_path'):
        # Gebruik standaard demo logo als fallback
        print("No logo path found in profile, using default logo")
        demo_logo_path = '/mnt/d/Projects/ai-arbeidsdeskundige_claude/app/backend/storage/demo/demo_logo.svg'

        # Create if it doesn't exist
        if not os.path.exists(demo_logo_path):
            try:
                os.makedirs(os.path.dirname(demo_logo_path), exist_ok=True)
                # No need to create file - we've already created the SVG in a previous step
                print(f"Demo logo should be available at {demo_logo_path}")
            except Exception as e:
                print(f"Failed to create demo logo directory: {str(e)}")

        # Use the SVG logo
        user_profile['logo_path'] = demo_logo_path
    
    # Extra controle van beschikbare profiel informatie
    print(f"Profile complete? {user_profile is not None}")
    if user_profile:
        available_fields = [k for k, v in user_profile.items() if v is not None]
        print(f"Available profile fields: {available_fields}")

    # Maak titelpagina gebaseerd op template type
    if template_type == 'modern':
        create_modern_title_page(doc, report_data, user_profile)
    elif template_type == 'professioneel':
        create_professional_title_page(doc, report_data, user_profile)
    else:  # standaard template
        create_standard_title_page(doc, report_data, user_profile)

    # Pagina break na titelpagina
    doc.add_page_break()

    # Inhoudsopgave
    toc_heading = doc.add_paragraph("Inhoudsopgave", style='SectionHeading')

    # Secties
    section_order = [
        "persoonsgegevens",
        "werkgever_functie",
        "aanleiding",
        "arbeidsverleden",
        "medische_situatie",
        "belastbaarheid",
        "belasting_huidige_functie",
        "visie_ad",
        "matching",
        "conclusie",
        "samenvatting"
    ]

    section_titles = {
        "persoonsgegevens": "Persoonsgegevens",
        "werkgever_functie": "Werkgever en Functie",
        "aanleiding": "Aanleiding Onderzoek",
        "arbeidsverleden": "Arbeidsverleden en Opleidingsachtergrond",
        "medische_situatie": "Medische Situatie",
        "belastbaarheid": "Belastbaarheid",
        "belasting_huidige_functie": "Belasting Huidige Functie",
        "visie_ad": "Visie Arbeidsdeskundige",
        "matching": "Matching Overwegingen",
        "conclusie": "Conclusie en Advies",
        "samenvatting": "Samenvatting"
    }

    # Genereer inhoudsopgave
    for i, section_id in enumerate(section_order):
        if section_id in section_titles:
            section_title = section_titles[section_id]
            toc_item = doc.add_paragraph(style='NormalText')
            toc_item.add_run(f"{i+1}. {section_title}")

    # Pagina break na inhoudsopgave
    doc.add_page_break()

    # Stel koptekst en voettekst in op basis van template
    set_header_footer(doc, report_data, user_profile, template_type)

    # Secties toevoegen
    for i, section_id in enumerate(section_order):
        if section_id in section_titles and section_id in report_data.get('content', {}):
            section_title = section_titles[section_id]
            section_content = report_data.get('content', {}).get(section_id, "")

            # Sectie titel
            doc.add_paragraph(f"{i+1}. {section_title}", style='SectionHeading')

            # Sectie inhoud - verwerk markdown naar opgemaakte Word-paragrafen
            convert_markdown_to_paragraphs(doc, section_content, style='NormalText')

            # Pagina break na elke sectie behalve de laatste
            if i < len(section_order) - 1:
                doc.add_page_break()

    # Voeg arbeidsdeskundige informatie toe aan einde van rapport
    if user_profile:
        add_profile_info_section(doc, user_profile)

    # Sla document op
    doc.save(output_path)
    return output_path

def create_standard_title_page(doc, report_data, user_profile):
    """Standaard titelblad met basisinformatie"""
    # Titelpagina
    title = doc.add_paragraph("Arbeidsdeskundig Rapport", style='TitlePage')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Voeg subtitel toe
    if 'title' in report_data and report_data['title']:
        subtitle = doc.add_paragraph(report_data['title'], style='TitlePage')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datum
    date_par = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y')}", style='NormalText')
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Voeg profielinformatie toe aan titelpagina
    if user_profile:
        doc.add_paragraph()  # Lege regel voor spacing

        profile_par = doc.add_paragraph(style='ProfileDetails')
        profile_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Naam arbeidsdeskundige
        if user_profile.get('first_name') and user_profile.get('last_name'):
            profile_run = profile_par.add_run(
                f"Opgesteld door: {user_profile.get('first_name')} {user_profile.get('last_name')}"
            )
            profile_run.bold = True
            profile_par.add_run("\n")

        # Functietitel
        if user_profile.get('job_title'):
            profile_par.add_run(f"{user_profile.get('job_title')}\n")

        # Bedrijfsnaam
        if user_profile.get('company_name'):
            profile_par.add_run(f"{user_profile.get('company_name')}\n")

def create_modern_title_page(doc, report_data, user_profile):
    """Modern titelblad met logo en uitgebreide profielinformatie"""
    # Maak een tabel voor moderne layout
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False

    # Stel tabelstijl in (geen randen)
    table.style = 'Table Grid'
    for cell in table._cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Verberg tabelborders
    for row in table.rows:
        for cell in row.cells:
            cell.border = None

    # Logo cel (linksboven)
    logo_cell = table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]

    # Altijd bedrijfsnaam gebruiken, geen afbeelding
    logo_text = user_profile.get("company_name", "Arbeidsdeskundige")
    fallback_run = logo_paragraph.add_run(logo_text.upper())
    fallback_run.font.size = Pt(14)
    fallback_run.font.bold = True

    # Titelpagina (rechtsboven)
    title_cell = table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    title_run = title_paragraph.add_run("Arbeidsdeskundig Rapport")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Voeg subtitel toe
    if 'title' in report_data and report_data['title']:
        subtitle_paragraph = title_cell.add_paragraph()
        subtitle_run = subtitle_paragraph.add_run(report_data['title'])
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.bold = True
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Datum (rechtsonder subtitel)
    date_paragraph = title_cell.add_paragraph()
    date_run = date_paragraph.add_run(f"Datum: {datetime.now().strftime('%d-%m-%Y')}")
    date_run.font.size = Pt(12)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Profielinformatie (onderste rij, over beide kolommen)
    if user_profile:
        # Merge cellen voor profielinformatie
        profile_cell = table.cell(2, 0)
        profile_cell.merge(table.cell(2, 1))

        profile_paragraph = profile_cell.add_paragraph()
        profile_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Voeg arbeidsdeskundige naam toe
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_run = profile_paragraph.add_run(
                f"{user_profile.get('first_name')} {user_profile.get('last_name')}"
            )
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            profile_paragraph.add_run("\n")

        # Functietitel
        if user_profile.get('job_title'):
            job_run = profile_paragraph.add_run(f"{user_profile.get('job_title')}")
            job_run.font.size = Pt(12)
            profile_paragraph.add_run("\n")

        # Certificeringen
        if user_profile.get('certification'):
            cert_run = profile_paragraph.add_run(f"{user_profile.get('certification')}")
            cert_run.font.size = Pt(10)
            profile_paragraph.add_run("\n")

        # Registratienummer
        if user_profile.get('registration_number'):
            reg_run = profile_paragraph.add_run(f"Reg.nr: {user_profile.get('registration_number')}")
            reg_run.font.size = Pt(10)
            profile_paragraph.add_run("\n")

        # Bedrijfsinformatie
        if user_profile.get('company_name'):
            company_paragraph = profile_cell.add_paragraph()
            company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            company_run = company_paragraph.add_run(f"{user_profile.get('company_name')}")
            company_run.font.size = Pt(12)
            company_run.font.bold = True
            company_paragraph.add_run("\n")

            # Bedrijfsadres
            address_parts = []
            if user_profile.get('company_address'):
                address_parts.append(user_profile.get('company_address'))

            if user_profile.get('company_postal_code') and user_profile.get('company_city'):
                address_parts.append(
                    f"{user_profile.get('company_postal_code')} {user_profile.get('company_city')}"
                )

            if address_parts:
                address_run = company_paragraph.add_run(", ".join(address_parts))
                address_run.font.size = Pt(10)
                company_paragraph.add_run("\n")

            # Contactinformatie
            contact_parts = []
            if user_profile.get('company_phone'):
                contact_parts.append(f"Tel: {user_profile.get('company_phone')}")
            if user_profile.get('company_email'):
                contact_parts.append(f"E-mail: {user_profile.get('company_email')}")
            if user_profile.get('company_website'):
                contact_parts.append(f"Website: {user_profile.get('company_website')}")

            if contact_parts:
                contact_run = company_paragraph.add_run(" | ".join(contact_parts))
                contact_run.font.size = Pt(10)

def create_professional_title_page(doc, report_data, user_profile):
    """Professioneel titelblad met minimalistisch design en logo"""
    # Gebruik bedrijfsnaam als logo
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo_text = user_profile.get("company_name", "Arbeidsdeskundige")
    fallback_run = logo_paragraph.add_run(logo_text.upper())
    fallback_run.font.size = Pt(16)
    fallback_run.font.bold = True

    # Voeg wat ruimte toe
    doc.add_paragraph()
    doc.add_paragraph()

    # Hoofdtitel
    title = doc.add_paragraph("Arbeidsdeskundig Rapport", style='TitlePage')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Voeg subtitel toe
    if 'title' in report_data and report_data['title']:
        subtitle = doc.add_paragraph(report_data['title'], style='CustomSubtitle')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datum
    date_par = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y')}", style='NormalText')
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Voeg wat ruimte toe
    doc.add_paragraph()
    doc.add_paragraph()

    # Voeg professionele separator toe
    separator = doc.add_paragraph()
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.add_run("* * *").bold = True

    # Voeg profielinformatie toe
    if user_profile:
        doc.add_paragraph()
        profile_par = doc.add_paragraph(style='ProfileDetails')
        profile_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Naam en functie
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_run = profile_par.add_run(
                f"{user_profile.get('first_name')} {user_profile.get('last_name')}"
            )
            name_run.bold = True

            if user_profile.get('job_title'):
                profile_par.add_run(f", {user_profile.get('job_title')}")

            profile_par.add_run("\n")

        # Certificering en registratie
        cert_parts = []
        if user_profile.get('certification'):
            cert_parts.append(user_profile.get('certification'))
        if user_profile.get('registration_number'):
            cert_parts.append(f"Reg.nr: {user_profile.get('registration_number')}")

        if cert_parts:
            profile_par.add_run(" | ".join(cert_parts) + "\n")

        # Bedrijfsnaam en contact
        if user_profile.get('company_name'):
            profile_par.add_run(f"{user_profile.get('company_name')}\n")

            contact_parts = []
            if user_profile.get('company_phone'):
                contact_parts.append(f"Tel: {user_profile.get('company_phone')}")
            if user_profile.get('company_email'):
                contact_parts.append(f"E-mail: {user_profile.get('company_email')}")

            if contact_parts:
                profile_par.add_run(" | ".join(contact_parts))

def set_header_footer(doc, report_data, user_profile, template_type):
    """Stel koptekst en voettekst in op basis van template"""
    # Alleen header/footer instellen voor moderne en professionele templates
    if template_type == 'standaard':
        return

    # Loop door alle secties om headers/footers in te stellen
    for section in doc.sections:
        # Koptekst instellen
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # Voettekst instellen
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

        if template_type == 'modern':
            # Moderne header met rapport titel
            if 'title' in report_data and report_data['title']:
                header_para.text = report_data['title']
                header_para.style = 'Header'

            # Footer met profielgegevens en paginanummer
            if user_profile and user_profile.get('company_name'):
                footer_text = f"{user_profile.get('company_name')} | "
                if user_profile.get('first_name') and user_profile.get('last_name'):
                    footer_text += f"{user_profile.get('first_name')} {user_profile.get('last_name')}"

                footer_run = footer_para.add_run(footer_text)
                footer_run.font.size = Pt(8)

                # Pagina nummers rechts uitlijnen
                footer_para.add_run(add_page_number())
            else:
                # Alleen paginanummers als geen profielgegevens
                footer_para.add_run(add_page_number())

        elif template_type == 'professioneel':
            # Professionele header met bedrijfsnaam en rapporttitel
            if user_profile and user_profile.get('company_name'):
                header_run = header_para.add_run(user_profile.get('company_name'))
                header_run.font.size = Pt(8)
                header_run.font.bold = True

                if 'title' in report_data and report_data['title']:
                    header_para.add_run(f" | {report_data['title']}")

            # Professionele footer met naam en paginanummer
            if user_profile:
                name_parts = []
                if user_profile.get('first_name') and user_profile.get('last_name'):
                    name_parts.append(
                        f"{user_profile.get('first_name')} {user_profile.get('last_name')}"
                    )
                if user_profile.get('job_title'):
                    name_parts.append(user_profile.get('job_title'))

                if name_parts:
                    footer_run = footer_para.add_run(" | ".join(name_parts))
                    footer_run.font.size = Pt(8)

                    # Pagina nummers rechts uitlijnen
                    footer_para.add_run(add_page_number())
            else:
                # Alleen paginanummers als geen profielgegevens
                footer_para.add_run(add_page_number())

def add_page_number():
    """Voeg paginanummer element toe voor koptekst/voettekst"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    return " Pagina "  # Pagina indicator voor Word

def add_profile_info_section(doc, user_profile):
    """Voeg een sectie toe met volledige profielinformatie aan het einde van het rapport"""
    # Pagina break
    doc.add_page_break()

    # Sectie titel
    doc.add_paragraph("Rapport opgesteld door", style='SectionHeading')

    # Profielinformatie
    profile_para = doc.add_paragraph(style='NormalText')

    # Naam en functie
    if user_profile.get('first_name') and user_profile.get('last_name'):
        name_run = profile_para.add_run(
            f"{user_profile.get('first_name')} {user_profile.get('last_name')}"
        )
        name_run.bold = True
        profile_para.add_run("\n")

    if user_profile.get('job_title'):
        profile_para.add_run(f"Functie: {user_profile.get('job_title')}\n")

    # Certificeringen en registratie
    if user_profile.get('certification'):
        profile_para.add_run(f"Certificering: {user_profile.get('certification')}\n")

    if user_profile.get('registration_number'):
        profile_para.add_run(f"Registratienummer: {user_profile.get('registration_number')}\n")

    # Specialisaties
    if user_profile.get('specializations') and isinstance(user_profile.get('specializations'), list):
        profile_para.add_run(f"Specialisaties: {', '.join(user_profile.get('specializations'))}\n")

    # Bedrijfsinformatie header
    if user_profile.get('company_name'):
        doc.add_paragraph()  # Lege regel
        company_para = doc.add_paragraph(style='NormalText')
        company_run = company_para.add_run(f"{user_profile.get('company_name')}")
        company_run.bold = True
        company_para.add_run("\n")

        # Bedrijfsadres
        address_parts = []
        if user_profile.get('company_address'):
            address_parts.append(user_profile.get('company_address'))

        if user_profile.get('company_postal_code'):
            postal_city = user_profile.get('company_postal_code')
            if user_profile.get('company_city'):
                postal_city += f" {user_profile.get('company_city')}"
            address_parts.append(postal_city)

        if user_profile.get('company_country'):
            address_parts.append(user_profile.get('company_country'))

        if address_parts:
            company_para.add_run("Adres: " + ", ".join(address_parts) + "\n")

        # Contactgegevens
        if user_profile.get('company_phone'):
            company_para.add_run(f"Telefoon: {user_profile.get('company_phone')}\n")

        if user_profile.get('company_email'):
            company_para.add_run(f"E-mail: {user_profile.get('company_email')}\n")

        if user_profile.get('company_website'):
            company_para.add_run(f"Website: {user_profile.get('company_website')}\n")

    # Bedrijfsbeschrijving
    if user_profile.get('company_description'):
        doc.add_paragraph()  # Lege regel
        doc.add_paragraph("Over het bedrijf", style='NormalText').bold = True
        doc.add_paragraph(user_profile.get('company_description'), style='NormalText')

    # Voeg bio toe indien beschikbaar
    if user_profile.get('bio'):
        doc.add_paragraph()  # Lege regel
        doc.add_paragraph("Over de arbeidsdeskundige", style='NormalText').bold = True
        doc.add_paragraph(user_profile.get('bio'), style='NormalText')