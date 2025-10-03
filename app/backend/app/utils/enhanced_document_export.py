"""
Professional Document Export System for AI-Arbeidsdeskundige
Provides comprehensive Word and PDF export functionality with professional templates
"""
import os
import io
import re
import json
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import base64

# Word document dependencies
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table

# PDF export dependencies
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as RLTable, TableStyle, PageBreak
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# HTML to PDF dependencies  
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# Template engine
try:
    from jinja2 import Template, Environment, FileSystemLoader
    JINJA2_AVAILABLE = True
except ImportError:
    JINJA2_AVAILABLE = False

# Import structured output components
try:
    from app.utils.structured_output_generator import SectionContent, ContentType, OutputFormatter
    STRUCTURED_OUTPUT_AVAILABLE = True
except ImportError:
    STRUCTURED_OUTPUT_AVAILABLE = False


class ExportTemplateManager:
    """Manages export templates for different document formats"""
    
    def __init__(self):
        self.templates = {
            'standaard': StandardTemplate(),
            'modern': ModernTemplate(), 
            'professioneel': ProfessionalTemplate(),
            'compact': CompactTemplate()
        }
    
    def get_template(self, template_type: str) -> 'BaseTemplate':
        """Get template instance by type"""
        return self.templates.get(template_type, self.templates['standaard'])


class BaseTemplate:
    """Base template class for export formats"""
    
    def __init__(self):
        self.name = "Base"
        self.description = "Base template"
        self.colors = {
            'primary': RGBColor(30, 64, 175),     # #1e40af
            'secondary': RGBColor(37, 99, 235),   # #2563eb  
            'accent': RGBColor(59, 130, 246),     # #3b82f6
            'text': RGBColor(0, 0, 0),            # Black
            'light_gray': RGBColor(156, 163, 175)  # Gray for subtle elements
        }
        self.fonts = {
            'title': {'name': 'Times New Roman', 'size': 18, 'bold': True},
            'heading1': {'name': 'Times New Roman', 'size': 16, 'bold': True},
            'heading2': {'name': 'Times New Roman', 'size': 14, 'bold': True}, 
            'heading3': {'name': 'Times New Roman', 'size': 12, 'bold': True},
            'body': {'name': 'Times New Roman', 'size': 11, 'bold': False},
            'caption': {'name': 'Times New Roman', 'size': 10, 'bold': False}
        }
    
    def apply_document_settings(self, doc: Document):
        """Apply template-specific document settings"""
        # Set page margins and size
        for section in doc.sections:
            section.page_width = Cm(21)    # A4 width
            section.page_height = Cm(29.7) # A4 height
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
    
    def create_title_page(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Create title page - to be implemented by subclasses"""
        pass
    
    def apply_styles(self, doc: Document):
        """Apply template-specific styles"""
        styles = doc.styles
        
        # Create or update styles
        self._create_style(styles, 'TitleStyle', WD_STYLE_TYPE.PARAGRAPH, 
                          font_config=self.fonts['title'], color=self.colors['primary'])
        self._create_style(styles, 'Heading1Style', WD_STYLE_TYPE.PARAGRAPH,
                          font_config=self.fonts['heading1'], color=self.colors['primary'])
        self._create_style(styles, 'Heading2Style', WD_STYLE_TYPE.PARAGRAPH,
                          font_config=self.fonts['heading2'], color=self.colors['secondary'])
        self._create_style(styles, 'Heading3Style', WD_STYLE_TYPE.PARAGRAPH, 
                          font_config=self.fonts['heading3'], color=self.colors['accent'])
        self._create_style(styles, 'BodyStyle', WD_STYLE_TYPE.PARAGRAPH,
                          font_config=self.fonts['body'], color=self.colors['text'])
    
    def _create_style(self, styles, style_name: str, style_type, font_config: Dict, color: RGBColor):
        """Helper to create or update a style"""
        try:
            style = styles.add_style(style_name, style_type)
        except ValueError:
            style = styles[style_name]
        
        font = style.font
        font.name = font_config['name']
        font.size = Pt(font_config['size'])
        font.bold = font_config['bold']
        font.color.rgb = color
    
    def set_headers_footers(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Set headers and footers - to be implemented by subclasses"""
        pass


class StandardTemplate(BaseTemplate):
    """Standard Dutch business template"""
    
    def __init__(self):
        super().__init__()
        self.name = "Standaard"
        self.description = "Traditional Dutch business format"
    
    def create_title_page(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Create standard title page"""
        # Title
        title = doc.add_paragraph("Arbeidsdeskundig Rapport", style='TitleStyle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if report_data.get('title'):
            subtitle = doc.add_paragraph(report_data['title'], style='Heading1Style')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y')}", style='BodyStyle')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Profile information
        if user_profile:
            doc.add_paragraph()  # Space
            self._add_profile_section(doc, user_profile)
    
    def _add_profile_section(self, doc: Document, user_profile: Dict):
        """Add professional profile section"""
        profile_para = doc.add_paragraph(style='BodyStyle')
        profile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_run = profile_para.add_run(
                f"Opgesteld door: {user_profile['first_name']} {user_profile['last_name']}"
            )
            name_run.bold = True
            profile_para.add_run("\n")
        
        if user_profile.get('job_title'):
            profile_para.add_run(f"{user_profile['job_title']}\n")
        
        if user_profile.get('company_name'):
            profile_para.add_run(f"{user_profile['company_name']}\n")


class ModernTemplate(BaseTemplate):
    """Modern professional template with enhanced styling"""
    
    def __init__(self):
        super().__init__()
        self.name = "Modern"
        self.description = "Modern professional layout with headers/footers"
    
    def create_title_page(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Create modern title page with professional layout"""
        # Create table for layout
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Company logo/name (top left)
        logo_cell = table.cell(0, 0)
        logo_para = logo_cell.paragraphs[0]
        logo_text = user_profile.get("company_name", "Arbeidsdeskundige") if user_profile else "Arbeidsdeskundige"
        logo_run = logo_para.add_run(logo_text.upper())
        logo_run.font.size = Pt(14)
        logo_run.font.bold = True
        logo_run.font.color.rgb = self.colors['primary']
        
        # Title (top right)
        title_cell = table.cell(0, 1)
        title_para = title_cell.paragraphs[0]
        title_run = title_para.add_run("Arbeidsdeskundig Rapport")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = self.colors['primary']
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Subtitle and date
        if report_data.get('title'):
            subtitle_para = title_cell.add_paragraph()
            subtitle_run = subtitle_para.add_run(report_data['title'])
            subtitle_run.font.size = Pt(16)
            subtitle_run.font.bold = True
            subtitle_run.font.color.rgb = self.colors['secondary']
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        date_para = title_cell.add_paragraph()
        date_run = date_para.add_run(f"Datum: {datetime.now().strftime('%d-%m-%Y')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Profile information (bottom merged cells)
        if user_profile:
            profile_cell = table.cell(2, 0)
            profile_cell.merge(table.cell(2, 1))
            self._add_comprehensive_profile(profile_cell, user_profile)
    
    def _add_comprehensive_profile(self, cell, user_profile: Dict):
        """Add comprehensive profile information"""
        if not user_profile:
            return
            
        profile_para = cell.paragraphs[0]
        profile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name and title
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_run = profile_para.add_run(
                f"{user_profile['first_name']} {user_profile['last_name']}"
            )
            name_run.font.size = Pt(14)
            name_run.font.bold = True
            name_run.font.color.rgb = self.colors['primary']
            profile_para.add_run("\n")
        
        if user_profile.get('job_title'):
            job_run = profile_para.add_run(user_profile['job_title'])
            job_run.font.size = Pt(12)
            profile_para.add_run("\n")
        
        # Certifications
        cert_parts = []
        if user_profile.get('certification'):
            cert_parts.append(user_profile['certification'])
        if user_profile.get('registration_number'):
            cert_parts.append(f"Reg.nr: {user_profile['registration_number']}")
        
        if cert_parts:
            cert_run = profile_para.add_run(" | ".join(cert_parts))
            cert_run.font.size = Pt(10)
            profile_para.add_run("\n")
        
        # Company information  
        if user_profile.get('company_name'):
            company_para = cell.add_paragraph()
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            company_run = company_para.add_run(user_profile['company_name'])
            company_run.font.size = Pt(12)
            company_run.font.bold = True
            company_run.font.color.rgb = self.colors['secondary']
            company_para.add_run("\n")
            
            # Contact info
            contact_parts = []
            if user_profile.get('company_phone'):
                contact_parts.append(f"Tel: {user_profile['company_phone']}")
            if user_profile.get('company_email'):
                contact_parts.append(f"E-mail: {user_profile['company_email']}")
            if user_profile.get('company_website'):
                contact_parts.append(f"Website: {user_profile['company_website']}")
            
            if contact_parts:
                contact_run = company_para.add_run(" | ".join(contact_parts))
                contact_run.font.size = Pt(10)
    
    def set_headers_footers(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Set modern headers and footers"""
        for section in doc.sections:
            # Header
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            
            if report_data.get('title'):
                header_run = header_para.add_run(report_data['title'])
                header_run.font.size = Pt(10)
                header_run.font.color.rgb = self.colors['primary']
            
            # Footer
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            
            footer_text = ""
            if user_profile and user_profile.get('company_name'):
                footer_text += user_profile['company_name']
                if user_profile.get('first_name') and user_profile.get('last_name'):
                    footer_text += f" | {user_profile['first_name']} {user_profile['last_name']}"
            
            if footer_text:
                footer_run = footer_para.add_run(footer_text)
                footer_run.font.size = Pt(8)
                footer_run.font.color.rgb = self.colors['light_gray']


class ProfessionalTemplate(BaseTemplate):
    """Professional minimalist template"""
    
    def __init__(self):
        super().__init__()
        self.name = "Professioneel"
        self.description = "Minimalist professional design"
    
    def create_title_page(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Create professional minimalist title page"""
        # Company name as header
        if user_profile and user_profile.get('company_name'):
            company_para = doc.add_paragraph()
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            company_run = company_para.add_run(user_profile['company_name'].upper())
            company_run.font.size = Pt(16)
            company_run.font.bold = True
            company_run.font.color.rgb = self.colors['primary']
        
        # Spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Main title
        title = doc.add_paragraph("Arbeidsdeskundig Rapport", style='TitleStyle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        if report_data.get('title'):
            subtitle = doc.add_paragraph(report_data['title'], style='Heading1Style')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y')}", style='BodyStyle')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional separator
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run("* * *")
        sep_run.bold = True
        sep_run.font.color.rgb = self.colors['accent']
        
        # Condensed profile info
        if user_profile:
            doc.add_paragraph()
            self._add_minimal_profile(doc, user_profile)
    
    def _add_minimal_profile(self, doc: Document, user_profile: Dict):
        """Add minimal professional profile"""
        profile_para = doc.add_paragraph(style='BodyStyle')
        profile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Name and title on one line
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_text = f"{user_profile['first_name']} {user_profile['last_name']}"
            if user_profile.get('job_title'):
                name_text += f", {user_profile['job_title']}"
            
            name_run = profile_para.add_run(name_text)
            name_run.bold = True
            name_run.font.color.rgb = self.colors['primary']
            profile_para.add_run("\n")
        
        # Credentials and company
        cred_parts = []
        if user_profile.get('certification'):
            cred_parts.append(user_profile['certification'])
        if user_profile.get('registration_number'):
            cred_parts.append(f"Reg.nr: {user_profile['registration_number']}")
        
        if cred_parts:
            profile_para.add_run(" | ".join(cred_parts) + "\n")
        
        if user_profile.get('company_name'):
            company_run = profile_para.add_run(user_profile['company_name'])
            company_run.font.color.rgb = self.colors['secondary']
            profile_para.add_run("\n")
        
        # Essential contact
        contact_parts = []
        if user_profile.get('company_phone'):
            contact_parts.append(f"Tel: {user_profile['company_phone']}")
        if user_profile.get('company_email'):
            contact_parts.append(f"E-mail: {user_profile['company_email']}")
        
        if contact_parts:
            contact_run = profile_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = self.colors['light_gray']


class CompactTemplate(BaseTemplate):
    """Compact template for space-efficient reports"""
    
    def __init__(self):
        super().__init__()
        self.name = "Compact"
        self.description = "Space-efficient format with smaller headers"
        # Smaller font sizes for compact layout
        self.fonts = {
            'title': {'name': 'Times New Roman', 'size': 16, 'bold': True},
            'heading1': {'name': 'Times New Roman', 'size': 14, 'bold': True},
            'heading2': {'name': 'Times New Roman', 'size': 12, 'bold': True},
            'heading3': {'name': 'Times New Roman', 'size': 11, 'bold': True},
            'body': {'name': 'Times New Roman', 'size': 10, 'bold': False},
            'caption': {'name': 'Times New Roman', 'size': 9, 'bold': False}
        }
    
    def apply_document_settings(self, doc: Document):
        """Apply compact document settings"""
        super().apply_document_settings(doc)
        # Smaller margins for compact layout
        for section in doc.sections:
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
    
    def create_title_page(self, doc: Document, report_data: Dict, user_profile: Dict):
        """Create compact title page"""
        # Single line header with company and report type
        header_para = doc.add_paragraph(style='TitleStyle')
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        company_name = user_profile.get('company_name', 'Arbeidsdeskundige') if user_profile else 'Arbeidsdeskundige'
        header_run = header_para.add_run(f"{company_name} - Arbeidsdeskundig Rapport")
        header_run.font.color.rgb = self.colors['primary']
        
        # Subtitle and date on same line
        if report_data.get('title'):
            subtitle_para = doc.add_paragraph(style='Heading1Style')
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(
                f"{report_data['title']} | {datetime.now().strftime('%d-%m-%Y')}"
            )
            subtitle_run.font.color.rgb = self.colors['secondary']
        
        # Minimal profile in single line
        if user_profile:
            profile_para = doc.add_paragraph(style='BodyStyle')
            profile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            profile_parts = []
            if user_profile.get('first_name') and user_profile.get('last_name'):
                profile_parts.append(f"{user_profile['first_name']} {user_profile['last_name']}")
            if user_profile.get('job_title'):
                profile_parts.append(user_profile['job_title'])
            if user_profile.get('registration_number'):
                profile_parts.append(f"Reg: {user_profile['registration_number']}")
            
            if profile_parts:
                profile_run = profile_para.add_run(" | ".join(profile_parts))
                profile_run.font.size = Pt(9)
                profile_run.font.color.rgb = self.colors['light_gray']
        
        # Separator line
        separator = doc.add_paragraph()
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep_run = separator.add_run("─" * 60)
        sep_run.font.color.rgb = self.colors['accent']


class EnhancedWordExporter:
    """Enhanced Word document exporter with professional templates"""
    
    def __init__(self, template_manager: ExportTemplateManager = None):
        self.template_manager = template_manager or ExportTemplateManager()
    
    def export_report(self, report_data: Dict, output_path: str, template_type: str = 'standaard') -> str:
        """
        Export report to Word document with professional formatting
        
        Args:
            report_data: Complete report data including content and metadata
            output_path: Path where the DOCX file should be saved
            template_type: Template to use (standaard, modern, professioneel, compact)
            
        Returns:
            Path to the generated file
        """
        # Get template
        template = self.template_manager.get_template(template_type)
        
        # Create document
        doc = Document()
        
        # Apply template settings and styles
        template.apply_document_settings(doc)
        template.apply_styles(doc)
        
        # Get user profile and metadata
        user_profile = self._get_user_profile(report_data)
        
        # Create title page
        template.create_title_page(doc, report_data, user_profile)
        
        # Page break after title page
        doc.add_page_break()
        
        # Create table of contents
        self._create_table_of_contents(doc, report_data)
        
        # Page break after TOC
        doc.add_page_break()
        
        # Set headers and footers
        template.set_headers_footers(doc, report_data, user_profile)
        
        # Add report sections
        self._add_report_sections(doc, report_data, template)
        
        # Add profile information section
        if user_profile:
            self._add_profile_information_section(doc, user_profile, template)
        
        # Save document
        doc.save(output_path)
        return output_path
    
    def _extract_sections_from_markdown(self, markdown_content: str) -> Dict[str, str]:
        """Extract sections from markdown content"""
        # Simple implementation - split by major headers and map to section keys
        sections = {}
        
        # Define mapping from markdown headers to section keys
        header_mapping = {
            'Samenvatting': 'samenvatting',
            'Vraagstelling': 'vraagstelling', 
            'Ondernomen activiteiten': 'ondernomen_activiteiten',
            'Voorgeschiedenis': 'voorgeschiedenis',
            'Gegevens werkgever': 'gegevens_werkgever',
            'Gegevens werknemer': 'gegevens_werknemer',
            'Belastbaarheid': 'belastbaarheid',
            'Eigen functie': 'eigen_functie',
            'Geschiktheid': 'geschiktheid_analyse',
            'Conclusie': 'conclusie',
            'Trajectplan': 'trajectplan',
            'Vervolg': 'vervolg'
        }
        
        # Split markdown by headers and assign to sections
        import re
        header_pattern = r'^#{1,3}\s+(.+)$'
        
        current_section = None
        current_content = []
        
        for line in markdown_content.split('\n'):
            header_match = re.match(header_pattern, line)
            if header_match:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                header_text = header_match.group(1).strip()
                current_section = None
                current_content = []
                
                # Find matching section key
                for header, section_key in header_mapping.items():
                    if header.lower() in header_text.lower():
                        current_section = section_key
                        break
            else:
                if current_section:
                    current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
            
        return sections
    
    def _get_user_profile(self, report_data: Dict) -> Optional[Dict]:
        """Extract user profile from report data"""
        # Try multiple locations where user profile might be stored
        profile_locations = [
            report_data.get('user_profile'),
            report_data.get('metadata', {}).get('user_profile'),
            report_data.get('report_metadata', {}).get('user_profile')
        ]
        
        for profile in profile_locations:
            if profile and isinstance(profile, dict):
                return profile
        
        return None
    
    def _create_table_of_contents(self, doc: Document, report_data: Dict):
        """Create professional table of contents"""
        toc_title = doc.add_paragraph("Inhoudsopgave", style='Heading1Style')
        
        # Section order and titles
        section_order = [
            ("samenvatting", "Samenvatting"),
            ("vraagstelling", "Vraagstelling"),
            ("ondernomen_activiteiten", "Ondernomen activiteiten"),
            ("voorgeschiedenis", "Voorgeschiedenis"),
            ("gegevens_werkgever", "Gegevens werkgever"),
            ("gegevens_werknemer", "Gegevens werknemer"),
            ("belastbaarheid", "Belastbaarheid van werknemer"),
            ("eigen_functie", "Eigen functie werknemer"),
            ("geschiktheid_analyse", "Geschiktheid analyse"),
            ("conclusie", "Conclusie"),
            ("trajectplan", "Trajectplan"),
            ("vervolg", "Vervolg")
        ]
        
        # Use same content extraction logic as sections
        content = {}
        
        if 'content' in report_data and isinstance(report_data['content'], dict):
            content_data = report_data['content']
            
            if 'sections' in content_data:
                content = content_data['sections']
            elif 'markdown' in content_data:
                content = self._extract_sections_from_markdown(content_data['markdown'])
            else:
                content = content_data
        else:
            content = report_data.get('content', {})
        
        for i, (section_id, section_title) in enumerate(section_order):
            if section_id in content:
                toc_item = doc.add_paragraph(style='BodyStyle')
                toc_run = toc_item.add_run(f"{i+1}. {section_title}")
                # Add tab stops for page numbers in future enhancement
    
    def _add_report_sections(self, doc: Document, report_data: Dict, template: BaseTemplate):
        """Add all report sections with proper formatting"""
        # Try different content structures
        content = {}
        
        # First try structured content format
        if 'content' in report_data and isinstance(report_data['content'], dict):
            content_data = report_data['content']
            
            # Check if it has structured format (with sections key)
            if 'sections' in content_data:
                content = content_data['sections']
            # Check if it has markdown/html format
            elif 'markdown' in content_data:
                # Parse markdown content to extract sections
                content = self._extract_sections_from_markdown(content_data['markdown'])
            # Use direct content keys
            else:
                content = content_data
        else:
            content = report_data.get('content', {})
        
        section_order = [
            ("samenvatting", "Samenvatting"),
            ("vraagstelling", "Vraagstelling"),
            ("ondernomen_activiteiten", "Ondernomen activiteiten"),
            ("voorgeschiedenis", "Voorgeschiedenis"),
            ("gegevens_werkgever", "Gegevens werkgever"),
            ("gegevens_werknemer", "Gegevens werknemer"),
            ("belastbaarheid", "Belastbaarheid van werknemer"),
            ("eigen_functie", "Eigen functie werknemer"),
            ("geschiktheid_analyse", "Geschiktheid analyse"),
            ("conclusie", "Conclusie"),
            ("trajectplan", "Trajectplan"),
            ("vervolg", "Vervolg")
        ]
        
        # Debug: print content keys to understand structure
        print(f"Content keys available for export: {list(content.keys())}")
        print(f"Looking for sections: {[s[0] for s in section_order]}")
        
        for i, (section_id, section_title) in enumerate(section_order):
            if section_id in content:
                # Section heading
                heading = doc.add_paragraph(f"{i+1}. {section_title}", style='Heading1Style')
                
                # Section content
                section_content = content[section_id]
                
                # Check if we have structured content
                if STRUCTURED_OUTPUT_AVAILABLE and self._is_structured_content(report_data, section_id):
                    self._add_structured_section(doc, report_data, section_id, template)
                else:
                    self._add_text_content(doc, section_content)
                
                # Page break after each section except the last
                if i < len([s for s in section_order if s[0] in content]) - 1:
                    doc.add_page_break()
    
    def _is_structured_content(self, report_data: Dict, section_id: str) -> bool:
        """Check if section has structured content available"""
        metadata = report_data.get('metadata', {})
        structured_content = metadata.get('structured_content', {})
        return section_id in structured_content
    
    def _add_structured_section(self, doc: Document, report_data: Dict, section_id: str, template: BaseTemplate):
        """Add section with structured content formatting"""
        if not STRUCTURED_OUTPUT_AVAILABLE:
            # Fallback to text content
            content = report_data.get('content', {}).get(section_id, '')
            self._add_text_content(doc, content)
            return
        
        # Get structured content
        metadata = report_data.get('metadata', {})
        structured_data = metadata.get('structured_content', {}).get(section_id, {})
        
        try:
            # Create SectionContent object
            section_content = SectionContent(**structured_data)
            
            # Add summary if available
            if section_content.summary:
                summary_para = doc.add_paragraph(style='BodyStyle')
                summary_run = summary_para.add_run("Samenvatting: ")
                summary_run.bold = True
                summary_para.add_run(section_content.summary)
                doc.add_paragraph()  # Space
            
            # Process main content elements
            for element in section_content.main_content:
                if element.type == ContentType.PARAGRAPH:
                    doc.add_paragraph(element.content, style='BodyStyle')
                
                elif element.type == ContentType.LIST:
                    self._add_structured_list(doc, element)
                
                elif element.type == ContentType.ASSESSMENT:
                    self._add_assessment_table(doc, element, template)
                
                elif element.type == ContentType.TABLE:
                    self._add_structured_table(doc, element)
                
                # Add spacing between elements
                doc.add_paragraph()
            
            # Add conclusions
            if section_content.conclusions:
                conclusions_para = doc.add_paragraph(style='Heading3Style')
                conclusions_para.add_run("Conclusies")
                
                for conclusion in section_content.conclusions:
                    conclusion_para = doc.add_paragraph(style='BodyStyle')
                    conclusion_para.add_run(f"• {conclusion}")
                
                doc.add_paragraph()  # Space
            
            # Add recommendations
            if section_content.recommendations:
                rec_para = doc.add_paragraph(style='Heading3Style')
                rec_para.add_run("Aanbevelingen")
                
                for rec in section_content.recommendations:
                    rec_para = doc.add_paragraph(style='BodyStyle')
                    priority_indicator = self._get_priority_indicator(rec.priority)
                    rec_para.add_run(f"{priority_indicator} {rec.action}")
                    if rec.rationale:
                        rec_para.add_run(f" ({rec.rationale})")
        
        except Exception as e:
            print(f"Error processing structured content for {section_id}: {e}")
            # Fallback to text content
            content = report_data.get('content', {}).get(section_id, '')
            self._add_text_content(doc, content)
    
    def _add_structured_list(self, doc: Document, element):
        """Add a structured list element"""
        from app.utils.structured_output_generator import ListItem
        
        for item in element.content:
            if isinstance(item, ListItem):
                list_para = doc.add_paragraph(style='BodyStyle')
                if item.label:
                    label_run = list_para.add_run(f"{item.label} ")
                    label_run.bold = True
                list_para.add_run(item.value)
                if item.detail:
                    detail_run = list_para.add_run(f" ({item.detail})")
                    detail_run.italic = True
    
    def _add_assessment_table(self, doc: Document, element, template: BaseTemplate):
        """Add assessment matrix as a professional table"""
        from app.utils.structured_output_generator import AssessmentItem
        
        if not element.content:
            return
        
        # Create table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Aspect', 'Capaciteit', 'Frequentie', 'Beperking']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = template.colors['primary']
        
        # Data rows
        for item in element.content:
            if isinstance(item, AssessmentItem):
                row_cells = table.add_row().cells
                row_cells[0].text = item.aspect
                row_cells[1].text = item.capacity
                row_cells[2].text = item.frequency or '-'
                row_cells[3].text = item.limitation or '-'
    
    def _add_structured_table(self, doc: Document, element):
        """Add a structured table element"""
        if not element.content:
            return
        
        # Assuming content is a list of rows
        rows_data = element.content
        if not rows_data:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
        table.style = 'Table Grid'
        
        # Fill table
        for i, row_data in enumerate(rows_data):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = str(cell_data)
    
    def _get_priority_indicator(self, priority) -> str:
        """Get visual indicator for priority level"""
        priority_indicators = {
            'critical': '🔴',
            'high': '🟠', 
            'medium': '🟡',
            'low': '🟢'
        }
        return priority_indicators.get(priority.value if hasattr(priority, 'value') else priority, '•')
    
    def _add_text_content(self, doc: Document, content: str):
        """Add text content with basic markdown processing"""
        if not content:
            doc.add_paragraph("Geen inhoud beschikbaar.", style='BodyStyle')
            return
        
        # Split content into lines and process
        lines = content.split('\n')
        current_paragraph_lines = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - end current paragraph
                if current_paragraph_lines:
                    para_text = ' '.join(current_paragraph_lines)
                    self._add_formatted_paragraph(doc, para_text)
                    current_paragraph_lines = []
                doc.add_paragraph()  # Add space
                continue
            
            # Check for headers (markdown style)
            if line.startswith('#'):
                # End current paragraph first
                if current_paragraph_lines:
                    para_text = ' '.join(current_paragraph_lines)
                    self._add_formatted_paragraph(doc, para_text)
                    current_paragraph_lines = []
                
                # Process header
                header_level = 0
                while header_level < len(line) and line[header_level] == '#':
                    header_level += 1
                
                header_text = line[header_level:].strip()
                
                if header_level == 1:
                    style = 'Heading1Style'
                elif header_level == 2:
                    style = 'Heading2Style'
                else:
                    style = 'Heading3Style'
                
                doc.add_paragraph(header_text, style=style)
                continue
            
            # Check for list items
            if line.startswith(('- ', '* ', '• ')):
                # End current paragraph first
                if current_paragraph_lines:
                    para_text = ' '.join(current_paragraph_lines)
                    self._add_formatted_paragraph(doc, para_text)
                    current_paragraph_lines = []
                
                # Add list item
                list_text = line[2:].strip()
                list_para = doc.add_paragraph(list_text, style='BodyStyle')
                list_para.paragraph_format.left_indent = Inches(0.25)
                continue
            
            # Regular line - add to current paragraph
            current_paragraph_lines.append(line)
        
        # Add final paragraph if exists
        if current_paragraph_lines:
            para_text = ' '.join(current_paragraph_lines)
            self._add_formatted_paragraph(doc, para_text)
    
    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add paragraph with basic text formatting"""
        para = doc.add_paragraph(style='BodyStyle')
        
        # Process basic markdown formatting
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                para.add_run(part)
    
    def _add_profile_information_section(self, doc: Document, user_profile: Dict, template: BaseTemplate):
        """Add comprehensive profile information section"""
        # Page break
        doc.add_page_break()
        
        # Section title
        title_para = doc.add_paragraph("Rapport opgesteld door", style='Heading1Style')
        
        # Personal information
        if user_profile.get('first_name') and user_profile.get('last_name'):
            name_para = doc.add_paragraph(style='BodyStyle')
            name_run = name_para.add_run(
                f"{user_profile['first_name']} {user_profile['last_name']}"
            )
            name_run.bold = True
            name_run.font.size = Pt(12)
            name_run.font.color.rgb = template.colors['primary']
            
            if user_profile.get('job_title'):
                name_para.add_run(f"\n{user_profile['job_title']}")
        
        # Professional information
        prof_info = []
        if user_profile.get('certification'):
            prof_info.append(f"Certificering: {user_profile['certification']}")
        if user_profile.get('registration_number'):
            prof_info.append(f"Registratienummer: {user_profile['registration_number']}")
        if user_profile.get('specializations'):
            if isinstance(user_profile['specializations'], list):
                prof_info.append(f"Specialisaties: {', '.join(user_profile['specializations'])}")
        
        if prof_info:
            for info in prof_info:
                doc.add_paragraph(info, style='BodyStyle')
        
        # Company information
        if user_profile.get('company_name'):
            doc.add_paragraph()  # Space
            company_para = doc.add_paragraph(style='BodyStyle')
            company_run = company_para.add_run(user_profile['company_name'])
            company_run.bold = True
            company_run.font.size = Pt(12)
            company_run.font.color.rgb = template.colors['secondary']
            
            # Address
            address_parts = []
            if user_profile.get('company_address'):
                address_parts.append(user_profile['company_address'])
            if user_profile.get('company_postal_code'):
                postal_city = user_profile['company_postal_code']
                if user_profile.get('company_city'):
                    postal_city += f" {user_profile['company_city']}"
                address_parts.append(postal_city)
            if user_profile.get('company_country'):
                address_parts.append(user_profile['company_country'])
            
            if address_parts:
                doc.add_paragraph(f"Adres: {', '.join(address_parts)}", style='BodyStyle')
            
            # Contact information
            contact_info = []
            if user_profile.get('company_phone'):
                contact_info.append(f"Telefoon: {user_profile['company_phone']}")
            if user_profile.get('company_email'):
                contact_info.append(f"E-mail: {user_profile['company_email']}")
            if user_profile.get('company_website'):
                contact_info.append(f"Website: {user_profile['company_website']}")
            
            for contact in contact_info:
                doc.add_paragraph(contact, style='BodyStyle')
        
        # Company description
        if user_profile.get('company_description'):
            doc.add_paragraph()  # Space
            desc_title = doc.add_paragraph("Over het bedrijf", style='Heading3Style')
            doc.add_paragraph(user_profile['company_description'], style='BodyStyle')
        
        # Professional bio
        if user_profile.get('bio'):
            doc.add_paragraph()  # Space
            bio_title = doc.add_paragraph("Over de arbeidsdeskundige", style='Heading3Style')
            doc.add_paragraph(user_profile['bio'], style='BodyStyle')


class PDFExporter:
    """Professional PDF exporter using ReportLab"""
    
    def __init__(self):
        if not PDF_AVAILABLE:
            raise ImportError("PDF export requires reportlab package")
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom PDF styles matching Word templates"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=18,
            textColor=colors.Color(30/255, 64/255, 175/255),  # Primary blue
            spaceAfter=20,
            alignment=TA_CENTER
        ))
        
        # Heading styles
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=colors.Color(30/255, 64/255, 175/255),
            spaceBefore=16,
            spaceAfter=12
        ))
        
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.Color(37/255, 99/255, 235/255),
            spaceBefore=12,
            spaceAfter=8
        ))
        
        # Body text style
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            fontName='Times-Roman',
            spaceBefore=6,
            spaceAfter=6,
            alignment=TA_JUSTIFY
        ))
    
    def export_report(self, report_data: Dict, output_path: str, template_type: str = 'standaard') -> str:
        """
        Export report to PDF with professional formatting
        
        Args:
            report_data: Complete report data
            output_path: Path where PDF should be saved
            template_type: Template type (affects styling)
            
        Returns:
            Path to generated PDF
        """
        # Create PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            rightMargin=2.5*cm,
            leftMargin=2.5*cm,
            topMargin=2.5*cm,
            bottomMargin=2.5*cm
        )
        
        # Build story (content)
        story = []
        
        # Add title page
        self._add_title_page(story, report_data, template_type)
        
        # Add page break
        story.append(PageBreak())
        
        # Add table of contents
        self._add_table_of_contents(story, report_data)
        
        # Add page break
        story.append(PageBreak())
        
        # Add report sections
        self._add_report_sections(story, report_data)
        
        # Add profile section
        user_profile = self._get_user_profile(report_data)
        if user_profile:
            story.append(PageBreak())
            self._add_profile_section(story, user_profile)
        
        # Build PDF
        doc.build(story)
        return output_path
    
    def _get_user_profile(self, report_data: Dict) -> Optional[Dict]:
        """Extract user profile from report data"""
        profile_locations = [
            report_data.get('user_profile'),
            report_data.get('metadata', {}).get('user_profile'),
            report_data.get('report_metadata', {}).get('user_profile')
        ]
        
        for profile in profile_locations:
            if profile and isinstance(profile, dict):
                return profile
        return None
    
    def _add_title_page(self, story: List, report_data: Dict, template_type: str):
        """Add title page to PDF"""
        user_profile = self._get_user_profile(report_data)
        
        # Company name/logo
        if user_profile and user_profile.get('company_name'):
            story.append(Paragraph(
                user_profile['company_name'].upper(),
                self.styles['CustomHeading1']
            ))
        
        story.append(Spacer(1, 20))
        
        # Main title
        story.append(Paragraph("Arbeidsdeskundig Rapport", self.styles['CustomTitle']))
        
        # Subtitle
        if report_data.get('title'):
            story.append(Paragraph(report_data['title'], self.styles['CustomHeading1']))
        
        # Date
        story.append(Paragraph(
            f"Datum: {datetime.now().strftime('%d-%m-%Y')}",
            self.styles['CustomBody']
        ))
        
        story.append(Spacer(1, 40))
        
        # Profile information
        if user_profile:
            profile_info = []
            if user_profile.get('first_name') and user_profile.get('last_name'):
                profile_info.append(f"<b>Opgesteld door:</b> {user_profile['first_name']} {user_profile['last_name']}")
            if user_profile.get('job_title'):
                profile_info.append(user_profile['job_title'])
            if user_profile.get('company_name'):
                profile_info.append(user_profile['company_name'])
            
            for info in profile_info:
                story.append(Paragraph(info, self.styles['CustomBody']))
    
    def _add_table_of_contents(self, story: List, report_data: Dict):
        """Add table of contents"""
        story.append(Paragraph("Inhoudsopgave", self.styles['CustomHeading1']))
        story.append(Spacer(1, 20))
        
        section_order = [
            ("samenvatting", "Samenvatting"),
            ("vraagstelling", "Vraagstelling"),
            ("ondernomen_activiteiten", "Ondernomen activiteiten"),
            ("voorgeschiedenis", "Voorgeschiedenis"),
            ("gegevens_werkgever", "Gegevens werkgever"),
            ("gegevens_werknemer", "Gegevens werknemer"),
            ("belastbaarheid", "Belastbaarheid van werknemer"),
            ("eigen_functie", "Eigen functie werknemer"),
            ("geschiktheid_analyse", "Geschiktheid analyse"),
            ("conclusie", "Conclusie"),
            ("trajectplan", "Trajectplan"),
            ("vervolg", "Vervolg")
        ]
        
        # Use same content extraction logic as sections
        content = {}
        
        if 'content' in report_data and isinstance(report_data['content'], dict):
            content_data = report_data['content']
            
            if 'sections' in content_data:
                content = content_data['sections']
            elif 'markdown' in content_data:
                content = self._extract_sections_from_markdown(content_data['markdown'])
            else:
                content = content_data
        else:
            content = report_data.get('content', {})
        
        for i, (section_id, section_title) in enumerate(section_order):
            if section_id in content:
                story.append(Paragraph(
                    f"{i+1}. {section_title}",
                    self.styles['CustomBody']
                ))
    
    def _add_report_sections(self, story: List, report_data: Dict):
        """Add all report sections"""
        # Try different content structures (same logic as DOCX export)
        content = {}
        
        if 'content' in report_data and isinstance(report_data['content'], dict):
            content_data = report_data['content']
            
            if 'sections' in content_data:
                content = content_data['sections']
            elif 'markdown' in content_data:
                content = self._extract_sections_from_markdown(content_data['markdown'])
            else:
                content = content_data
        else:
            content = report_data.get('content', {})
        
        section_order = [
            ("samenvatting", "Samenvatting"),
            ("vraagstelling", "Vraagstelling"),
            ("ondernomen_activiteiten", "Ondernomen activiteiten"),
            ("voorgeschiedenis", "Voorgeschiedenis"),
            ("gegevens_werkgever", "Gegevens werkgever"),
            ("gegevens_werknemer", "Gegevens werknemer"),
            ("belastbaarheid", "Belastbaarheid van werknemer"),
            ("eigen_functie", "Eigen functie werknemer"),
            ("geschiktheid_analyse", "Geschiktheid analyse"),
            ("conclusie", "Conclusie"),
            ("trajectplan", "Trajectplan"),
            ("vervolg", "Vervolg")
        ]
        
        for i, (section_id, section_title) in enumerate(section_order):
            if section_id in content:
                # Add section heading
                story.append(Paragraph(
                    f"{i+1}. {section_title}",
                    self.styles['CustomHeading1']
                ))
                
                # Add section content
                section_content = content[section_id]
                self._add_text_content_pdf(story, section_content)
                
                # Add spacing
                story.append(Spacer(1, 20))
    
    def _add_text_content_pdf(self, story: List, content: str):
        """Add text content to PDF with formatting"""
        if not content:
            story.append(Paragraph("Geen inhoud beschikbaar.", self.styles['CustomBody']))
            return
        
        # Split content into paragraphs and process
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Process basic markdown
            formatted_text = self._process_markdown_pdf(para_text)
            story.append(Paragraph(formatted_text, self.styles['CustomBody']))
    
    def _process_markdown_pdf(self, text: str) -> str:
        """Process basic markdown for PDF"""
        # Bold text
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        # Italic text
        text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
        # Line breaks
        text = text.replace('\n', '<br/>')
        return text
    
    def _add_profile_section(self, story: List, user_profile: Dict):
        """Add profile section to PDF"""
        story.append(Paragraph("Rapport opgesteld door", self.styles['CustomHeading1']))
        
        # Personal info
        if user_profile.get('first_name') and user_profile.get('last_name'):
            story.append(Paragraph(
                f"<b>{user_profile['first_name']} {user_profile['last_name']}</b>",
                self.styles['CustomBody']
            ))
        
        if user_profile.get('job_title'):
            story.append(Paragraph(
                f"Functie: {user_profile['job_title']}",
                self.styles['CustomBody']
            ))
        
        # Add other profile information
        profile_fields = [
            ('certification', 'Certificering'),
            ('registration_number', 'Registratienummer'),
            ('company_name', 'Bedrijf'),
            ('company_phone', 'Telefoon'),
            ('company_email', 'E-mail'),
            ('company_website', 'Website')
        ]
        
        for field, label in profile_fields:
            if user_profile.get(field):
                story.append(Paragraph(
                    f"{label}: {user_profile[field]}",
                    self.styles['CustomBody']
                ))


# Main export functions that integrate with existing system
def export_report_to_docx_enhanced(
    report_data: Dict, 
    output_path: str, 
    template_type: str = 'standaard'
) -> str:
    """
    Enhanced Word export function with professional templates
    
    This function replaces the original export_report_to_docx function
    with enhanced professional formatting and template support
    """
    exporter = EnhancedWordExporter()
    return exporter.export_report(report_data, output_path, template_type)


def export_report_to_pdf(
    report_data: Dict, 
    output_path: str, 
    template_type: str = 'standaard'
) -> str:
    """
    Export report to PDF format with professional styling
    
    Args:
        report_data: Complete report data including content and metadata
        output_path: Path where the PDF file should be saved
        template_type: Template to use (affects styling)
        
    Returns:
        Path to the generated PDF file
    """
    if not PDF_AVAILABLE:
        raise ImportError("PDF export requires reportlab package. Install with: pip install reportlab")
    
    exporter = PDFExporter()
    return exporter.export_report(report_data, output_path, template_type)


def export_report_to_html(
    report_data: Dict, 
    output_path: str, 
    template_type: str = 'standaard'
) -> str:
    """
    Export report to HTML format for web viewing or HTML-to-PDF conversion
    
    Args:
        report_data: Complete report data
        output_path: Path where HTML should be saved
        template_type: Template type
        
    Returns:
        Path to generated HTML file
    """
    # Create HTML template
    html_template = """
    <!DOCTYPE html>
    <html lang="nl">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{{ title }}</title>
        <style>
            body {
                font-family: 'Times New Roman', serif;
                font-size: 11pt;
                line-height: 1.5;
                margin: 0;
                padding: 2.5cm;
                color: #000;
            }
            .title-page {
                text-align: center;
                margin-bottom: 2cm;
            }
            .title {
                font-size: 18pt;
                font-weight: bold;
                color: #1e40af;
                margin-bottom: 1cm;
            }
            .subtitle {
                font-size: 16pt;
                font-weight: bold;
                color: #2563eb;
                margin-bottom: 0.5cm;
            }
            .date {
                font-size: 11pt;
                margin-bottom: 1cm;
            }
            .profile-info {
                font-size: 11pt;
                margin-top: 2cm;
            }
            h1 {
                font-size: 16pt;
                font-weight: bold;
                color: #1e40af;
                margin-top: 1.5cm;
                margin-bottom: 0.5cm;
            }
            h2 {
                font-size: 14pt;
                font-weight: bold;
                color: #2563eb;
                margin-top: 1cm;
                margin-bottom: 0.3cm;
            }
            h3 {
                font-size: 12pt;
                font-weight: bold;
                color: #3b82f6;
                margin-top: 0.8cm;
                margin-bottom: 0.3cm;
            }
            p {
                margin-bottom: 0.5cm;
            }
            .toc {
                margin-bottom: 2cm;
            }
            .toc-item {
                margin-bottom: 0.2cm;
            }
            table {
                width: 100%;
                border-collapse: collapse;
                margin: 1cm 0;
            }
            table, th, td {
                border: 1px solid #ccc;
            }
            th, td {
                padding: 8pt;
                text-align: left;
            }
            th {
                background-color: #f8f9fa;
                font-weight: bold;
                color: #1e40af;
            }
            .page-break {
                page-break-before: always;
            }
        </style>
    </head>
    <body>
        <!-- Title Page -->
        <div class="title-page">
            {% if user_profile and user_profile.company_name %}
            <div style="font-size: 14pt; font-weight: bold; color: #1e40af;">
                {{ user_profile.company_name.upper() }}
            </div>
            <br><br>
            {% endif %}
            
            <div class="title">Arbeidsdeskundig Rapport</div>
            
            {% if title %}
            <div class="subtitle">{{ title }}</div>
            {% endif %}
            
            <div class="date">Datum: {{ date }}</div>
            
            {% if user_profile %}
            <div class="profile-info">
                {% if user_profile.first_name and user_profile.last_name %}
                <strong>Opgesteld door: {{ user_profile.first_name }} {{ user_profile.last_name }}</strong><br>
                {% endif %}
                {% if user_profile.job_title %}
                {{ user_profile.job_title }}<br>
                {% endif %}
                {% if user_profile.company_name %}
                {{ user_profile.company_name }}<br>
                {% endif %}
            </div>
            {% endif %}
        </div>
        
        <!-- Table of Contents -->
        <div class="page-break">
            <h1>Inhoudsopgave</h1>
            <div class="toc">
                {% for section_id, section_title in sections %}
                <div class="toc-item">{{ loop.index }}. {{ section_title }}</div>
                {% endfor %}
            </div>
        </div>
        
        <!-- Report Sections -->
        {% for section_id, section_title in sections %}
        <div class="page-break">
            <h1>{{ loop.index }}. {{ section_title }}</h1>
            {{ content[section_id] | safe }}
        </div>
        {% endfor %}
        
        <!-- Profile Section -->
        {% if user_profile %}
        <div class="page-break">
            <h1>Rapport opgesteld door</h1>
            
            {% if user_profile.first_name and user_profile.last_name %}
            <p><strong>{{ user_profile.first_name }} {{ user_profile.last_name }}</strong></p>
            {% endif %}
            
            {% if user_profile.job_title %}
            <p>Functie: {{ user_profile.job_title }}</p>
            {% endif %}
            
            {% if user_profile.certification %}
            <p>Certificering: {{ user_profile.certification }}</p>
            {% endif %}
            
            {% if user_profile.registration_number %}
            <p>Registratienummer: {{ user_profile.registration_number }}</p>
            {% endif %}
            
            {% if user_profile.company_name %}
            <h3>{{ user_profile.company_name }}</h3>
            
            {% if user_profile.company_address %}
            <p>Adres: {{ user_profile.company_address }}
            {% if user_profile.company_postal_code %}
            , {{ user_profile.company_postal_code }}
            {% endif %}
            {% if user_profile.company_city %}
            {{ user_profile.company_city }}
            {% endif %}
            </p>
            {% endif %}
            
            {% if user_profile.company_phone %}
            <p>Telefoon: {{ user_profile.company_phone }}</p>
            {% endif %}
            
            {% if user_profile.company_email %}
            <p>E-mail: {{ user_profile.company_email }}</p>
            {% endif %}
            
            {% if user_profile.company_website %}
            <p>Website: {{ user_profile.company_website }}</p>
            {% endif %}
            
            {% endif %}
        </div>
        {% endif %}
    </body>
    </html>
    """
    
    if not JINJA2_AVAILABLE:
        raise ImportError("HTML export requires jinja2 package. Install with: pip install jinja2")
    
    # Prepare template data
    user_profile = None
    profile_locations = [
        report_data.get('user_profile'),
        report_data.get('metadata', {}).get('user_profile'),
        report_data.get('report_metadata', {}).get('user_profile')
    ]
    
    for profile in profile_locations:
        if profile and isinstance(profile, dict):
            user_profile = profile
            break
    
    sections = [
        ("persoonsgegevens", "Persoonsgegevens"),
        ("werkgever_functie", "Werkgever en Functie"),
        ("aanleiding", "Aanleiding Onderzoek"),
        ("arbeidsverleden", "Arbeidsverleden en Opleidingsachtergrond"),
        ("medische_situatie", "Medische Situatie"),
        ("belastbaarheid", "Belastbaarheid"),
        ("belasting_huidige_functie", "Belasting Huidige Functie"),
        ("visie_ad", "Visie Arbeidsdeskundige"),
        ("matching", "Matching Overwegingen"),
        ("conclusie", "Conclusie en Advies"),
        ("samenvatting", "Samenvatting")
    ]
    
    # Filter sections that exist in content
    content = report_data.get('content', {})
    available_sections = [(sid, title) for sid, title in sections if sid in content]
    
    # Process content for HTML (convert markdown to HTML)
    html_content = {}
    for section_id, section_content in content.items():
        html_content[section_id] = _markdown_to_html(section_content)
    
    # Render template
    template = Template(html_template)
    html_output = template.render(
        title=report_data.get('title', ''),
        date=datetime.now().strftime('%d-%m-%Y'),
        user_profile=user_profile,
        sections=available_sections,
        content=html_content
    )
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_output)
    
    return output_path


def _markdown_to_html(text: str) -> str:
    """Convert basic markdown to HTML"""
    if not text:
        return ""
    
    # Basic markdown to HTML conversion
    html = text
    
    # Headers
    html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)  
    html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    
    # Bold and italic
    html = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html)
    
    # Lists
    html = re.sub(r'^\- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'(<li>.*</li>)', r'<ul>\1</ul>', html, flags=re.DOTALL)
    
    # Paragraphs
    paragraphs = html.split('\n\n')
    html_paragraphs = []
    for para in paragraphs:
        para = para.strip()
        if para and not para.startswith('<'):
            para = f'<p>{para}</p>'
        html_paragraphs.append(para)
    
    html = '\n'.join(html_paragraphs)
    
    # Line breaks
    html = html.replace('\n', '<br>\n')
    
    return html


# Compatibility function - maintains existing interface
def export_report_to_docx(report_data: Dict, output_path: str) -> str:
    """
    Compatibility function for existing code - uses enhanced export
    """
    # Get template type from report metadata
    template_type = 'standaard'  # Default
    
    # Try to get template type from metadata
    metadata = report_data.get('metadata', {}) or report_data.get('report_metadata', {})
    if metadata and 'template_type' in metadata:
        template_type = metadata['template_type']
    
    return export_report_to_docx_enhanced(report_data, output_path, template_type)