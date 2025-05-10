import os
import sys
import json

# Voorbeeld rapportstructuur voor een AD rapport in staatvandienst format
sample_report = {
    "persoonsgegevens": "Dit is informatie over de persoon zoals naam, geboortedatum, adres.",
    "werkgever_functie": "Informatie over de huidige/laatste werkgever en functie.",
    "aanleiding": "Reden voor het arbeidsdeskundig onderzoek.",
    "arbeidsverleden": "Overzicht van opleidingen en werkervaring.",
    "medische_situatie": "Beschrijving van de medische situatie.",
    "belastbaarheid": "Analyse van wat de persoon wel/niet kan.",
    "belasting_huidige_functie": "Analyse van de eisen van de huidige functie.",
    "visie_ad": "Professionele visie van de arbeidsdeskundige.",
    "matching": "Overwegingen voor matching naar passend werk.",
    "conclusie": "Conclusies en aanbevelingen.",
    "samenvatting": "Korte samenvatting van het rapport."
}

def generate_docx_export_function():
    """
    Genereer code voor een functie die een rapport exporteert naar DOCX formaat
    """
    code = """
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def export_report_to_docx(report_data, output_path):
    '''
    Exporteer rapport naar DOCX formaat
    
    Args:
        report_data (dict): Rapport data met alle secties
        output_path (str): Pad waar het DOCX bestand moet worden opgeslagen
    
    Returns:
        str: Pad naar het gegenereerde bestand
    '''
    # Maak nieuw document
    doc = Document()
    
    # Configureer document en stijlen
    sections = doc.sections
    for section in sections:
        section.page_width = Cm(21)  # A4 breedte
        section.page_height = Cm(29.7)  # A4 hoogte
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    
    # Stijlen definiëren
    styles = doc.styles
    
    # Titelpagina stijl
    title_style = styles.add_style('TitlePage', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(24)
    title_font.bold = True
    
    # Hoofdstuktitel stijl
    heading_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(14)
    heading_font.bold = True
    
    # Normale tekst stijl
    normal_style = styles.add_style('NormalText', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    
    # Titelpagina
    title = doc.add_paragraph("Arbeidsdeskundig Rapport", style='TitlePage')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Voeg subtitel toe
    if 'title' in report_data and report_data['title']:
        subtitle = doc.add_paragraph(report_data['title'], style='TitlePage')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datum
    date_par = doc.add_paragraph(f"Datum: {datetime.now().strftime('%d-%m-%Y')}", style='NormalText')
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pagina break na titelpagina
    doc.add_page_break()
    
    # Inhoudsopgave
    toc_heading = doc.add_paragraph("Inhoudsopgave", style='SectionHeading')
    
    # Secties
    section_order = [
        "persoonsgegevens",
        "werkgever_functie",
        "aanleiding",
        "arbeidsverleden",
        "medische_situatie",
        "belastbaarheid",
        "belasting_huidige_functie",
        "visie_ad",
        "matching",
        "conclusie",
        "samenvatting"
    ]
    
    section_titles = {
        "persoonsgegevens": "Persoonsgegevens",
        "werkgever_functie": "Werkgever en Functie",
        "aanleiding": "Aanleiding Onderzoek",
        "arbeidsverleden": "Arbeidsverleden en Opleidingsachtergrond",
        "medische_situatie": "Medische Situatie",
        "belastbaarheid": "Belastbaarheid",
        "belasting_huidige_functie": "Belasting Huidige Functie",
        "visie_ad": "Visie Arbeidsdeskundige",
        "matching": "Matching Overwegingen",
        "conclusie": "Conclusie en Advies",
        "samenvatting": "Samenvatting"
    }
    
    # Genereer inhoudsopgave
    for i, section_id in enumerate(section_order):
        if section_id in section_titles:
            section_title = section_titles[section_id]
            toc_item = doc.add_paragraph(style='NormalText')
            toc_item.add_run(f"{i+1}. {section_title}")
    
    # Pagina break na inhoudsopgave
    doc.add_page_break()
    
    # Secties toevoegen
    for i, section_id in enumerate(section_order):
        if section_id in section_titles and section_id in report_data.get('content', {}):
            section_title = section_titles[section_id]
            section_content = report_data.get('content', {}).get(section_id, "")
            
            # Sectie titel
            doc.add_paragraph(f"{i+1}. {section_title}", style='SectionHeading')
            
            # Sectie inhoud - split op regeleinden voor betere opmaak
            paragraphs = section_content.split('\\n')
            for p in paragraphs:
                if p.strip():
                    doc.add_paragraph(p.strip(), style='NormalText')
            
            # Pagina break na elke sectie behalve de laatste
            if i < len(section_order) - 1:
                doc.add_page_break()
    
    # Sla document op
    doc.save(output_path)
    return output_path
"""
    return code

def generate_api_endpoint():
    """
    Genereer code voor een API endpoint om rapporten te exporteren
    """
    code = """
@router.get("/{report_id}/export/docx", response_class=FileResponse)
async def export_report_docx(report_id: UUID, user_info = Depends(verify_token)):
    '''
    Export a report to DOCX format
    '''
    user_id = user_info["user_id"]
    
    try:
        # Check if report exists and belongs to user
        report = db_service.get_report(str(report_id), user_id)
        
        if not report:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Report not found"
            )
        
        # Create temporary directory if it doesn't exist
        temp_dir = os.path.join(settings.STORAGE_PATH, "temp")
        os.makedirs(temp_dir, exist_ok=True)
        
        # Generate a filename based on report title and ID
        safe_title = "".join(c for c in report["title"] if c.isalnum() or c in [' ', '_']).rstrip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{safe_title}_rapport_{str(report_id)[:8]}.docx"
        output_path = os.path.join(temp_dir, filename)
        
        # Generate the DOCX file
        from app.utils.document_export import export_report_to_docx
        file_path = export_report_to_docx(report, output_path)
        
        # Return the file as a download
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error exporting report: {str(e)}"
        )
"""
    return code

def main():
    print("Voorbeeld rapport structuur:")
    print(json.dumps(sample_report, indent=2))
    
    print("\nVoorheel export functie (docx):")
    print(generate_docx_export_function())
    
    print("\nVoorbeeld API endpoint:")
    print(generate_api_endpoint())

if __name__ == "__main__":
    main()